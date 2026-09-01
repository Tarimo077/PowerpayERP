import csv
import calendar
import mimetypes
import textwrap
from datetime import date, timedelta
from decimal import Decimal
from io import BytesIO
from django.conf import settings
from django.contrib import messages
from django.contrib.auth import login
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth.models import User
from django.db import transaction
from django.db.models import Count, Q, Sum
from django.http import FileResponse, HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.utils import timezone
from django.utils.timesince import timesince
from django.views.decorators.http import require_POST
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas
from .decorators import roles_allowed
from .forms import *
from .models import *
from .emailing import send_invite_email, send_leave_email, send_task_assignment_email


def _profile(request):
    return get_object_or_404(
        Profile.objects.select_related("organization", "department"), user=request.user
    )


def _reporting_descendant_ids(manager):
    relationships = list(
        Profile.objects.filter(
            organization=manager.organization, user__is_active=True
        ).values_list("id", "manager_id")
    )
    children = {}
    for profile_id, manager_id in relationships:
        children.setdefault(manager_id, []).append(profile_id)
    descendants = []
    queue = list(children.get(manager.id, []))
    seen = {manager.id}
    while queue:
        profile_id = queue.pop(0)
        if profile_id in seen:
            continue
        seen.add(profile_id)
        descendants.append(profile_id)
        queue.extend(children.get(profile_id, []))
    return descendants


def _audit(request, action, obj):
    p = getattr(request.user, "profile", None)
    details = {
        "method": request.method,
        "path": request.path,
        "label": str(obj),
    }
    status_value = getattr(obj, "status", None)
    if status_value is not None:
        details["status"] = status_value
    AuditLog.objects.create(
        organization=getattr(p, "organization", None),
        actor=request.user,
        action=action,
        entity_type=obj.__class__.__name__,
        entity_id=str(obj.pk),
        details=details,
        ip_address=request.META.get("REMOTE_ADDR"),
    )


def home(request):
    return (
        redirect("dashboard")
        if request.user.is_authenticated
        else render(request, "erp/landing.html")
    )


def register(request):
    if request.user.is_authenticated:
        return redirect("dashboard")
    form = RegistrationForm(request.POST or None)

    if request.method == "POST" and form.is_valid():
        form.save()
        messages.success(
            request,
            "Your organization workspace is ready. Sign in with your email to continue.",
        )
        return redirect("login")

    return render(request, "registration/register.html", {"form": form})


@login_required
def dashboard(request):
    if request.user.is_superuser:
        return render(
            request,
            "erp/dashboard.html",
            {
                "super_view": True,
                "stats": [
                    ("Organizations", Organization.objects.count()),
                    ("Active users", User.objects.filter(is_active=True).count()),
                    ("Tasks", Task.objects.count()),
                    ("Audit events", AuditLog.objects.count()),
                ],
                "recent_audit": AuditLog.objects.select_related(
                    "actor", "organization"
                )[:8],
            },
        )

    p = _profile(request)
    org = p.organization
    tasks = Task.objects.filter(organization=org)

    if p.role == "employee":
        tasks = tasks.filter(Q(assigned_to=p) | Q(created_by=request.user)).distinct()
    elif p.role == "manager":
        team_ids = _reporting_descendant_ids(p)
        tasks = tasks.filter(
            Q(assigned_to=p)
            | Q(created_by=request.user)
            | Q(assigned_to_id__in=team_ids)
        ).distinct()

    sheets = Timesheet.objects.filter(organization=org)

    if p.role == "employee":
        sheets = sheets.filter(employee=p)

    elif p.role == "manager":
        sheets = sheets.filter(
            Q(employee=p)
            | Q(
                employee_id__in=team_ids,
                status__in=["submitted", "approved", "rejected"],
            )
        )

    stats = [
        ("Active tasks", tasks.exclude(status__in=["completed", "approved"]).count()),
        (
            "Overdue",
            tasks.filter(due_date__lt=timezone.localdate())
            .exclude(status__in=["completed", "approved"])
            .count(),
        ),
        ("Pending approvals", sheets.filter(status="submitted").count()),
        (
            "Team members",
            (
                Profile.objects.filter(organization=org, user__is_active=True).count()
                if p.role == "admin"
                else len(team_ids) if p.role == "manager" else 0
            ),
        ),
    ]
    calendar_tasks = (
        (
            Task.objects.filter(organization=org, assigned_to=p)
            if p.role in ["employee", "manager"]
            else tasks
        )
        .exclude(status__in=["completed", "approved"])
        .select_related("assigned_to__user")
    )

    def event(task):
        return {
            "title": task.title,
            "start": task.start_date.isoformat(),
            "due": task.due_date.isoformat(),
            "priority": task.priority,
            "status": task.status,
            "status_label": task.get_status_display(),
            "url": reverse("task_detail", args=[task.pk]),
            "assignee": str(task.assigned_to),
        }

    task_events = [event(task) for task in calendar_tasks]
    team_task_events = []
    team_member_count = 0
    if p.role == "manager":
        team_ids = _reporting_descendant_ids(p)
        team_member_count = len(team_ids)
        team_tasks = (
            Task.objects.filter(organization=org, assigned_to_id__in=team_ids)
            .exclude(status__in=["completed", "approved"])
            .select_related("assigned_to__user")
        )
        team_task_events = [event(task) for task in team_tasks]
    return render(
        request,
        "erp/dashboard.html",
        {
            "stats": stats,
            "tasks": tasks.select_related("assigned_to__user")[:8],
            "timesheets": sheets.select_related("employee__user")[:6],
            "task_events": task_events,
            "team_task_events": team_task_events,
            "team_member_count": team_member_count,
        },
    )


@login_required
def profile_page(request):
    if request.user.is_superuser:
        messages.info(request, "Platform administrators do not have employee profiles.")
        return redirect("dashboard")

    profile = get_object_or_404(
        Profile.objects.select_related("organization", "department", "manager__user"),
        user=request.user,
    )
    form = ProfileNameForm(request.POST or None, user=request.user)
    if request.method == "POST" and form.is_valid():
        form.save()
        _audit(request, "profile_name_updated", profile)
        messages.success(request, "Your name has been updated.")
        return redirect("profile")

    return render(
        request,
        "erp/profile.html",
        {"employee_profile": profile, "form": form},
    )


@login_required
def storyboard(request):
    return render(request, "erp/storyboard.html")


@login_required
def global_search(request):
    profile = _profile(request)
    query = request.GET.get("q", "").strip()
    results = {
        "tasks": Task.objects.none(),
        "employees": Profile.objects.none(),
        "documents": Document.objects.none(),
        "vouchers": PaymentVoucher.objects.none(),
        "item_requests": ItemRequest.objects.none(),
    }
    if len(query) >= 2:
        tasks = Task.objects.filter(organization=profile.organization)
        if profile.role == "employee":
            tasks = tasks.filter(Q(assigned_to=profile) | Q(created_by=request.user))
        elif profile.role == "manager":
            tasks = tasks.filter(
                Q(assigned_to=profile)
                | Q(created_by=request.user)
                | Q(assigned_to_id__in=_reporting_descendant_ids(profile))
            )
        results["tasks"] = tasks.filter(
            Q(title__icontains=query) | Q(description__icontains=query)
        ).select_related("assigned_to__user")[:10]

        if profile.role in {"manager", "admin"}:
            employees = Profile.objects.filter(organization=profile.organization)
            if profile.role == "manager":
                employees = employees.filter(pk__in=_reporting_descendant_ids(profile))
            results["employees"] = employees.filter(
                Q(user__first_name__icontains=query)
                | Q(user__last_name__icontains=query)
                | Q(user__email__icontains=query)
                | Q(employee_id__icontains=query)
            ).select_related("user", "department")[:10]

        results["documents"] = _visible_documents(profile).filter(
            Q(title__icontains=query) | Q(category__icontains=query)
        )[:10]
        results["vouchers"] = _visible_vouchers(profile).filter(
            Q(number__icontains=query) | Q(payee__icontains=query)
        )[:10]
        results["item_requests"] = (
            _visible_item_requests(profile)
            .filter(
                Q(number__icontains=query)
                | Q(purpose__icontains=query)
                | Q(lines__item__icontains=query)
            )
            .distinct()[:10]
        )

    result_count = sum(result.count() for result in results.values())
    return render(
        request,
        "erp/search_results.html",
        {"query": query, "result_count": result_count, **results},
    )


def _chat_title(thread, viewer):
    if thread.kind == "group":
        return thread.name
    other = (
        thread.memberships.exclude(profile=viewer)
        .select_related("profile__user")
        .first()
    )
    return str(other.profile) if other else "Direct message"


def _presence_data(profile):
    presence = getattr(profile, "chat_presence", None)
    online = bool(presence and presence.is_online)
    return {
        "online": online,
        "status": (
            "Online"
            if online
            else (
                f"Last seen {timesince(presence.last_seen)} ago"
                if presence
                else "Not seen yet"
            )
        ),
    }


def _message_data(message, viewer):
    return {
        "id": message.pk,
        "body": message.body,
        "mine": message.sender_id == viewer.id,
        "sender": str(message.sender),
        "initials": (
            (message.sender.user.first_name[:1] + message.sender.user.last_name[:1])
            or message.sender.user.username[:2]
        ).upper(),
        "time": timezone.localtime(message.created_at).strftime("%H:%M"),
        "date": timezone.localtime(message.created_at).strftime("%d %b %Y"),
        "attachment": (
            reverse("chat_attachment", args=[message.pk]) if message.attachment else ""
        ),
        "attachment_name": (
            message.attachment.name.rsplit("/", 1)[-1] if message.attachment else ""
        ),
    }


@login_required
def chats(request):
    p = _profile(request)
    memberships = (
        ChatMembership.objects.filter(
            profile=p, thread__organization=p.organization, thread__is_active=True
        )
        .select_related("thread")
        .prefetch_related("thread__memberships__profile__user")
    )
    rows = []
    for membership in memberships:
        thread = membership.thread
        last = thread.messages.last()
        thread.chat_title = _chat_title(thread, p)
        thread.last_message = last
        thread.unread_count = (
            thread.messages.filter(created_at__gt=membership.last_read_at)
            .exclude(sender=p)
            .count()
        )
        if thread.kind == "direct":
            other = (
                thread.memberships.exclude(profile=p)
                .select_related("profile__user", "profile__chat_presence")
                .first()
            )
            thread.other_profile = other.profile if other else None
            thread.other_presence = (
                _presence_data(other.profile)
                if other
                else {"online": False, "status": "Unavailable"}
            )
        rows.append(thread)
    rows.sort(
        key=lambda thread: (
            thread.last_message.created_at if thread.last_message else thread.created_at
        ),
        reverse=True,
    )
    invitations = ChatInvitation.objects.filter(
        invitee=p, status="pending", thread__organization=p.organization
    ).select_related("thread", "invited_by")
    online_profiles = (
        Profile.objects.filter(
            organization=p.organization,
            user__is_active=True,
            chat_presence__last_seen__gte=timezone.now() - timedelta(minutes=2),
        )
        .exclude(pk=p.pk)
        .select_related("user", "chat_presence")[:12]
    )
    direct_threads = [thread for thread in rows if thread.kind == "direct"]
    group_threads = [thread for thread in rows if thread.kind == "group"]
    return render(
        request,
        "erp/chat_list.html",
        {
            "chat_threads": rows,
            "direct_threads": direct_threads,
            "group_threads": group_threads,
            "chat_invitations": invitations,
            "online_profiles": online_profiles,
        },
    )


@login_required
def direct_chat_create(request):
    p = _profile(request)
    form = DirectChatForm(
        request.POST or None, organization=p.organization, current_profile=p
    )
    if request.method == "POST" and form.is_valid():
        other = form.cleaned_data["person"]
        direct_key = "-".join(str(value) for value in sorted([p.pk, other.pk]))
        with transaction.atomic():
            thread, created = ChatThread.objects.get_or_create(
                organization=p.organization,
                direct_key=direct_key,
                defaults={"kind": "direct", "created_by": request.user},
            )
            for member in [p, other]:
                ChatMembership.objects.get_or_create(
                    thread=thread, profile=member, defaults={"added_by": request.user}
                )
        if created:
            _audit(request, "direct_chat_started", thread)
        return redirect("chat_detail", pk=thread.pk)
    return render(
        request,
        "erp/chat_form.html",
        {"form": form, "title": "Start a conversation", "submit": "Start chat"},
    )


@login_required
def group_chat_create(request):
    p = _profile(request)
    form = GroupChatForm(
        request.POST or None, organization=p.organization, current_profile=p
    )
    if request.method == "POST" and form.is_valid():
        with transaction.atomic():
            thread = form.save(False)
            thread.organization = p.organization
            thread.kind = "group"
            thread.created_by = request.user
            thread.full_clean()
            thread.save()
            ChatMembership.objects.create(
                thread=thread, profile=p, added_by=request.user, is_admin=True
            )
            for invitee in form.cleaned_data["invitees"]:
                invitation = ChatInvitation(
                    thread=thread, invitee=invitee, invited_by=request.user
                )
                invitation.full_clean()
                invitation.save()
                Notification.objects.create(
                    user=invitee.user,
                    title="Channel invitation",
                    message=f"{request.user.get_full_name() or request.user.username} invited you to {thread.name}.",
                    url="/chats/",
                )
            _audit(request, "chat_group_created", thread)
        messages.success(request, f"{thread.name} created. Invitations were sent.")
        return redirect("chat_detail", pk=thread.pk)
    return render(
        request,
        "erp/chat_form.html",
        {"form": form, "title": "Create channel or group", "submit": "Create group"},
    )


@login_required
def chat_detail(request, pk):
    p = _profile(request)
    membership = get_object_or_404(
        ChatMembership.objects.select_related("thread"),
        profile=p,
        thread_id=pk,
        thread__organization=p.organization,
        thread__is_active=True,
    )
    thread = membership.thread
    form = ChatMessageForm(request.POST or None, request.FILES or None)
    if request.method == "POST" and form.is_valid():
        message = form.save(False)
        message.thread = thread
        message.sender = p
        message.full_clean()
        message.save()
        ChatThread.objects.filter(pk=thread.pk).update(updated_at=timezone.now())
        membership.last_read_at = timezone.now()
        membership.save(update_fields=["last_read_at", "updated_at"])
        data = _message_data(message, p)
        if request.headers.get("X-Requested-With") == "XMLHttpRequest":
            return JsonResponse({"ok": True, "message": data})
        return redirect("chat_detail", pk=thread.pk)
    membership.last_read_at = timezone.now()
    membership.save(update_fields=["last_read_at", "updated_at"])
    thread.chat_title = _chat_title(thread, p)
    chat_messages_qs = list(
        thread.messages.select_related("sender__user").order_by("-created_at")[:100]
    )
    chat_messages_qs.reverse()
    members = []
    for item in thread.memberships.select_related(
        "profile__user", "profile__chat_presence"
    ).all():
        item.presence_data = _presence_data(item.profile)
        members.append(item)
    return render(
        request,
        "erp/chat_detail.html",
        {
            "thread": thread,
            "membership": membership,
            "chat_messages": chat_messages_qs,
            "chat_members": members,
            "form": form,
        },
    )


@login_required
def chat_message_feed(request, pk):
    p = _profile(request)
    membership = get_object_or_404(
        ChatMembership,
        profile=p,
        thread_id=pk,
        thread__organization=p.organization,
        thread__is_active=True,
    )
    try:
        after = max(0, int(request.GET.get("after", 0)))
    except ValueError:
        after = 0
    feed = ChatMessage.objects.filter(thread_id=pk, pk__gt=after).select_related(
        "sender__user"
    )[:100]
    membership.last_read_at = timezone.now()
    membership.save(update_fields=["last_read_at", "updated_at"])
    members = []
    for item in membership.thread.memberships.select_related(
        "profile__user", "profile__chat_presence"
    ):
        presence = _presence_data(item.profile)
        members.append(
            {
                "id": item.profile_id,
                "name": str(item.profile),
                "initials": (
                    (item.profile.user.first_name[:1] + item.profile.user.last_name[:1])
                    or item.profile.user.username[:2]
                ).upper(),
                **presence,
            }
        )
    return JsonResponse(
        {
            "messages": [_message_data(message, p) for message in feed],
            "members": members,
        }
    )


@login_required
def chat_attachment(request, pk):
    p = _profile(request)
    message = get_object_or_404(
        ChatMessage.objects.select_related("thread"),
        pk=pk,
        thread__organization=p.organization,
        thread__memberships__profile=p,
    )
    if not message.attachment:
        return HttpResponse(status=404)
    return FileResponse(
        message.attachment.open("rb"),
        as_attachment=True,
        filename=message.attachment.name.rsplit("/", 1)[-1],
    )


@login_required
def chat_invite(request, pk):
    p = _profile(request)
    membership = get_object_or_404(
        ChatMembership,
        thread_id=pk,
        profile=p,
        thread__organization=p.organization,
        is_admin=True,
    )
    thread = membership.thread
    if thread.kind != "group":
        messages.error(request, "Only groups can invite members.")
        return redirect("chat_detail", pk=pk)
    form = ChatInviteForm(request.POST or None, thread=thread)
    if request.method == "POST" and form.is_valid():
        for invitee in form.cleaned_data["people"]:
            invitation = ChatInvitation(
                thread=thread, invitee=invitee, invited_by=request.user
            )
            invitation.full_clean()
            invitation.save()
            Notification.objects.create(
                user=invitee.user,
                title="Channel invitation",
                message=f"You were invited to {thread.name}.",
                url="/chats/",
            )
        _audit(request, "chat_members_invited", thread)
        messages.success(request, "Channel invitations sent.")
        return redirect("chat_detail", pk=pk)
    return render(
        request,
        "erp/chat_form.html",
        {
            "form": form,
            "title": f"Invite to {thread.name}",
            "submit": "Send invitations",
        },
    )


@login_required
@require_POST
def chat_invitation_action(request, pk, action):
    p = _profile(request)
    invitation = get_object_or_404(
        ChatInvitation,
        pk=pk,
        invitee=p,
        status="pending",
        thread__organization=p.organization,
    )
    if action not in ["accept", "decline"]:
        messages.error(request, "That invitation action is not allowed.")
        return redirect("chats")
    invitation.status = "accepted" if action == "accept" else "declined"
    invitation.responded_at = timezone.now()
    invitation.save(update_fields=["status", "responded_at", "updated_at"])
    if action == "accept":
        ChatMembership.objects.get_or_create(
            thread=invitation.thread,
            profile=p,
            defaults={"added_by": invitation.invited_by},
        )
        messages.success(request, f"You joined {invitation.thread.name}.")
        return redirect("chat_detail", pk=invitation.thread_id)
    messages.info(request, f"Invitation to {invitation.thread.name} declined.")
    return redirect("chats")


@login_required
@require_POST
def chat_leave(request, pk):
    p = _profile(request)
    membership = get_object_or_404(
        ChatMembership, thread_id=pk, profile=p, thread__organization=p.organization
    )
    if membership.thread.kind != "group":
        messages.error(request, "Direct conversations cannot be left.")
        return redirect("chat_detail", pk=pk)
    thread = membership.thread
    was_admin = membership.is_admin
    membership.delete()
    if was_admin and not thread.memberships.filter(is_admin=True).exists():
        successor = thread.memberships.order_by("created_at").first()
        if successor:
            successor.is_admin = True
            successor.save(update_fields=["is_admin", "updated_at"])
    messages.info(request, f"You left {thread.name}.")
    return redirect("chats")


@login_required
def tasks(request):
    p = _profile(request)

    qs = Task.objects.filter(organization=p.organization).select_related(
        "assigned_to__user", "department"
    )

    if p.role == "employee":
        qs = qs.filter(Q(assigned_to=p) | Q(created_by=request.user)).distinct()

    elif p.role == "manager":
        qs = qs.filter(
            Q(assigned_to=p)
            | Q(created_by=request.user)
            | Q(department=p.department)
            | Q(assigned_to__manager=p)
        ).distinct()

    q = request.GET.get("q", "")
    status = request.GET.get("status", "")

    if q:
        qs = qs.filter(
            Q(title__icontains=q)
            | Q(description__icontains=q)
            | Q(assigned_to__user__first_name__icontains=q)
        )
    if status:
        qs = qs.filter(status=status)

    task_rows = list(qs)
    for task in task_rows:
        task.can_change_status = bool(_task_status_choices(task, p))
    return render(
        request, "erp/task_list.html", {"tasks": task_rows, "statuses": Task.STATUSES}
    )


@login_required
def task_create(request):
    p = _profile(request)
    task = Task(organization=p.organization, created_by=request.user)
    form = TaskForm(
        request.POST or None,
        request.FILES or None,
        organization=p.organization,
        current_profile=p,
        instance=task,
    )
    if request.method == "POST" and form.is_valid():
        obj = form.save(False)
        obj.organization = p.organization
        obj.created_by = request.user
        # A newly delegated task has not started until its assignee says so.
        obj.status = "assigned"
        obj.full_clean()
        obj.save()
        Notification.objects.create(
            user=obj.assigned_to.user,
            title="New task assigned",
            message=obj.title,
            url=f"/tasks/{obj.pk}/",
        )
        send_task_assignment_email(obj)
        _audit(request, "task_created", obj)
        messages.success(request, "Task assigned successfully.")
        return redirect("task_create" if request.POST.get("_add_another") else "tasks")

    return render(
        request,
        "erp/form.html",
        {"form": form, "title": "Assign task", "submit": "Create task"},
    )


@login_required
def task_detail(request, pk):
    p = _profile(request)
    obj = get_object_or_404(Task, pk=pk, organization=p.organization)
    if (
        p.role == "employee"
        and obj.assigned_to_id != p.id
        and obj.created_by_id != request.user.id
    ):
        return redirect("tasks")
    return render(
        request,
        "erp/task_detail.html",
        {"task": obj, "can_change_status": bool(_task_status_choices(obj, p))},
    )


@login_required
def task_attachment(request, pk):
    profile = _profile(request)
    queryset = Task.objects.filter(organization=profile.organization)
    if profile.role == "employee":
        queryset = queryset.filter(Q(assigned_to=profile) | Q(created_by=request.user))
    elif profile.role == "manager":
        queryset = queryset.filter(
            Q(assigned_to=profile)
            | Q(created_by=request.user)
            | Q(assigned_to_id__in=_reporting_descendant_ids(profile))
        )
    task = get_object_or_404(queryset.distinct(), pk=pk)
    if not task.attachment:
        return HttpResponse(status=404)
    return FileResponse(
        task.attachment.open("rb"),
        as_attachment=True,
        filename=task.attachment.name.rsplit("/", 1)[-1],
    )


@login_required
def task_status(request, pk):
    p = _profile(request)
    obj = get_object_or_404(Task, pk=pk, organization=p.organization)
    choices = _task_status_choices(obj, p)
    if not choices:
        messages.error(request, "You cannot change the status of this task.")
        return redirect("task_detail", pk=pk)
    form = TaskStatusForm(request.POST or None, choices=choices)
    if request.method == "POST" and form.is_valid():
        previous = obj.get_status_display()
        obj.status = form.cleaned_data["status"]
        obj.save(update_fields=["status", "updated_at"])
        note = form.cleaned_data["note"]
        if obj.assigned_to.user_id != request.user.id:
            Notification.objects.create(
                user=obj.assigned_to.user,
                title=f"Task status changed to {obj.get_status_display()}",
                message=note
                or f"{request.user.get_full_name() or request.user.username} changed {obj.title} from {previous} to {obj.get_status_display()}.",
                url=f"/tasks/{obj.pk}/",
            )
        _audit(request, "task_status_changed", obj)
        messages.success(request, f"Task status changed to {obj.get_status_display()}.")
        return redirect("task_detail", pk=pk)
    return render(
        request,
        "erp/task_status_form.html",
        {"form": form, "task": obj, "title": "Change task status"},
    )


def _task_status_choices(task, profile):
    if hasattr(task, "payment_voucher_approval") or hasattr(
        task, "item_request_approval"
    ):
        return []
    if profile.role == "employee":
        if task.assigned_to_id != profile.id:
            return []
        allowed = {
            "pending": ["in_progress", "submitted"],
            "assigned": ["in_progress", "submitted"],
            "in_progress": ["submitted"],
            "rejected": ["in_progress", "submitted"],
        }.get(task.status, [])
    elif profile.role == "manager":
        in_scope = (
            task.assigned_to_id == profile.id
            or task.created_by_id == profile.user_id
            or task.assigned_to.manager_id == profile.id
            or (profile.department_id and task.department_id == profile.department_id)
        )
        if not in_scope:
            return []
        allowed = [value for value, _ in Task.STATUSES if value != task.status]
    elif profile.role == "admin":
        allowed = [value for value, _ in Task.STATUSES if value != task.status]
    else:
        return []
    labels = dict(Task.STATUSES)
    if profile.role == "employee":
        labels.update(
            {"in_progress": "Start / resume work", "submitted": "Submit for review"}
        )
    return [(value, labels[value]) for value in allowed]


@login_required
def timesheets(request):
    p = _profile(request)
    qs = (
        Timesheet.objects.filter(organization=p.organization)
        .select_related("employee__user", "reviewed_by")
        .annotate(hours=Sum("entries__hours"))
    )
    if p.role == "employee":
        qs = qs.filter(employee=p)
    elif p.role == "manager":
        team_ids = _reporting_descendant_ids(p)
        qs = qs.filter(
            Q(employee=p)
            | (
                ~Q(status="draft")
                & (Q(employee_id__in=team_ids) | Q(requested_approver=p))
            )
        ).distinct()
    try:
        year = int(request.GET.get("year", timezone.localdate().year))
    except (TypeError, ValueError):
        year = timezone.localdate().year
    year = max(2000, min(2100, year))
    current_year = timezone.localdate().year
    recorded_years = set(
        qs.order_by().values_list("period_start__year", flat=True).distinct()
    )
    year_options = sorted(
        {current_year - 1, current_year, current_year + 1, year, *recorded_years}
    )
    qs = qs.filter(period_start__year=year)
    rows = list(qs)
    for sheet in rows:
        sheet.can_edit_draft = sheet.status == "draft" and _can_edit_timesheet(sheet, p)
        sheet.can_delete_draft = sheet.status == "draft" and _can_delete_timesheet(
            sheet, p
        )
    months = []
    for month in range(1, 13):
        months.append(
            {
                "number": month,
                "name": calendar.month_name[month],
                "short": calendar.month_abbr[month],
                "timesheets": [
                    sheet for sheet in rows if sheet.period_start.month == month
                ],
            }
        )
    export_anchor = next((sheet for sheet in rows if sheet.employee_id == p.id), None)
    return render(
        request,
        "erp/timesheet_list.html",
        {
            "timesheets": rows,
            "months": months,
            "year": year,
            "year_options": year_options,
            "export_anchor": export_anchor,
        },
    )


@login_required
def timesheet_create(request):
    p = _profile(request)
    form = TimesheetForm(
        request.POST or None,
        request.FILES or None,
        organization=p.organization,
        employee=p,
    )
    if request.method == "POST" and form.is_valid():
        obj = form.save(False)
        month = int(form.cleaned_data["month"])
        year = form.cleaned_data["year"]
        obj.period_start = date(year, month, 1)
        obj.period_end = date(year, month, calendar.monthrange(year, month)[1])
        obj.organization = p.organization
        obj.employee = p
        obj.save()
        assigned_tasks = Task.objects.filter(
            organization=p.organization,
            assigned_to=p,
            actual_started_at__date__lte=obj.period_end,
        ).filter(
            Q(actual_completed_at__isnull=True)
            | Q(actual_completed_at__date__gte=obj.period_start)
        )
        prefilled = _prefilled_timesheet_entries(obj, assigned_tasks)
        if prefilled:
            TimesheetEntry.objects.bulk_create(prefilled)
        _audit(request, "timesheet_created", obj)
        messages.success(
            request,
            f"Timesheet created with {len(prefilled)} assigned task{'s' if len(prefilled)!=1 else ''} prefilled. Add hours and adjust the rows before submitting.",
        )
        if request.POST.get("_add_another"):
            return redirect("timesheet_create")
        return redirect("timesheet_detail", pk=obj.pk)
    return render(
        request,
        "erp/form.html",
        {"form": form, "title": "New timesheet", "submit": "Create timesheet"},
    )


def _timesheet_task_activity(task):
    description = (task.description or "").strip()
    return f"{task.title}: {description}" if description else task.title


def _prefilled_timesheet_entries(timesheet, tasks):
    work_by_date = {}
    today = timezone.localdate()
    for task in tasks:
        if not task.actual_started_at:
            continue
        started = timezone.localdate(task.actual_started_at)
        finished = (
            timezone.localdate(task.actual_completed_at)
            if task.actual_completed_at
            else today
        )
        first = max(started, timesheet.period_start)
        last = min(finished, timesheet.period_end)
        if first > last:
            continue
        for offset in range((last - first).days + 1):
            work_date = first + timedelta(days=offset)
            if work_date.weekday() >= 5:
                continue
            work_by_date.setdefault(work_date, []).append(task)
    entries = []
    for work_date, active_tasks in work_by_date.items():
        task_count = Decimal(len(active_tasks))
        for task in active_tasks:
            entries.append(
                TimesheetEntry(
                    timesheet=timesheet,
                    date=work_date,
                    task=task,
                    task_performed=_timesheet_task_activity(task),
                    hours=Decimal("8") / task_count,
                    days_worked=Decimal("1") / task_count,
                    location=timesheet.country,
                    description=task.description,
                )
            )
    return entries


@login_required
def timesheet_edit(request, pk):
    p = _profile(request)
    obj = get_object_or_404(Timesheet, pk=pk, organization=p.organization)
    if obj.status != "draft" or not _can_edit_timesheet(obj, p):
        messages.error(request, "Only an accessible draft timesheet can be edited.")
        return redirect("timesheet_detail", pk=pk)
    form = TimesheetForm(
        request.POST or None,
        request.FILES or None,
        organization=p.organization,
        employee=obj.employee,
        instance=obj,
    )
    if request.method == "POST" and form.is_valid():
        updated = form.save(False)
        updated.organization = p.organization
        updated.employee = obj.employee
        updated.save()
        _audit(request, "timesheet_updated", updated)
        messages.success(request, "Draft timesheet updated.")
        return redirect("timesheet_detail", pk=pk)
    return render(
        request,
        "erp/form.html",
        {"form": form, "title": "Edit draft timesheet", "submit": "Save changes"},
    )


@login_required
@require_POST
def timesheet_delete(request, pk):
    p = _profile(request)
    obj = get_object_or_404(
        Timesheet.objects.select_related("employee__user", "request_task"),
        pk=pk,
        organization=p.organization,
    )
    if obj.status != "draft" or not _can_delete_timesheet(obj, p):
        messages.error(request, "Only an accessible draft timesheet can be deleted.")
        return redirect("timesheet_detail", pk=pk)
    label = f"{obj.period_start:%B %Y} timesheet for {obj.employee}"
    request_task = obj.request_task
    _audit(request, "timesheet_deleted", obj)
    if request_task and request_task.status not in ["approved", "completed"]:
        request_task.status = "completed"
        request_task.save(update_fields=["status", "updated_at"])
    employee_user = obj.employee.user
    actor_user_id = request.user.id
    obj.delete()
    if employee_user.id != actor_user_id:
        Notification.objects.create(
            user=employee_user,
            title="Draft timesheet deleted",
            message=label,
            url="/timesheets/",
        )
    messages.info(request, f"{label} was deleted.")
    return redirect("timesheets")


@login_required
def timesheet_detail(request, pk):
    p = _profile(request)
    obj = get_object_or_404(
        Timesheet.objects.prefetch_related("entries"),
        pk=pk,
        organization=p.organization,
    )
    if p.role == "employee" and obj.employee_id != p.id:
        return redirect("timesheets")
    if p.role == "manager" and not _manager_can_view_timesheet(obj, p):
        return redirect("timesheets")
    form = EntryForm(
        request.POST or None,
        request.FILES or None,
        organization=p.organization,
        employee=obj.employee,
    )

    can_edit = _can_edit_timesheet(obj, p)
    if request.method == "POST" and form.is_valid() and can_edit:
        entry = form.save(False)
        entry.timesheet = obj
        entry.full_clean()
        entry.save()
        _audit(request, "timesheet_entry_added", entry)
        messages.success(request, "Timesheet entry added.")
        return redirect("timesheet_detail", pk=pk)
    entries = list(obj.entries.select_related("task").order_by("date", "created_at"))
    by_date = {}
    for entry in entries:
        by_date.setdefault(entry.date, []).append(entry)
    daily_rows = []
    for day in range(1, obj.period_end.day + 1):
        current = obj.period_start.replace(day=day)
        items = by_date.get(current, [])
        daily_rows.append(
            {
                "date": current,
                "entries": items,
                "days_worked": sum((item.days_worked for item in items), 0),
                "hours": sum((item.hours for item in items), 0),
                "overnight_duty": any(item.overnight_duty_station for item in items),
                "overnight_outside": any(
                    item.overnight_outside_duty_station for item in items
                ),
                "per_diem": any(item.per_diem_requested for item in items),
                "locations": ", ".join(
                    dict.fromkeys(item.location for item in items if item.location)
                ),
                "activities": "\n".join(item.task_performed for item in items),
            }
        )
    capacity_conflicts = [
        row for row in daily_rows if row["hours"] > 8 or row["days_worked"] > 1
    ]
    cumulative = (
        TimesheetEntry.objects.filter(
            timesheet__organization=p.organization,
            timesheet__employee=obj.employee,
            timesheet__period_start__lte=obj.period_end,
        ).aggregate(total=Sum("days_worked"))["total"]
        or 0
    )
    remaining = (
        (obj.initial_budget_days - cumulative)
        if obj.initial_budget_days is not None
        else None
    )
    consumption = (
        (cumulative / obj.initial_budget_days * 100)
        if obj.initial_budget_days
        else None
    )
    return render(
        request,
        "erp/timesheet_detail.html",
        {
            "timesheet": obj,
            "form": form,
            "daily_rows": daily_rows,
            "capacity_conflicts": capacity_conflicts,
            "cumulative_days": cumulative,
            "remaining_days": remaining,
            "consumption_rate": consumption,
            "can_review": _can_review_timesheet(obj, p),
            "can_edit": can_edit,
            "can_edit_draft": obj.status == "draft" and can_edit,
            "can_delete_draft": obj.status == "draft" and _can_delete_timesheet(obj, p),
            "is_owner": obj.employee_id == p.id,
        },
    )


@login_required
def timesheet_entry_edit(request, pk, entry_pk):
    p = _profile(request)
    timesheet = get_object_or_404(Timesheet, pk=pk, organization=p.organization)
    if not _can_edit_timesheet(timesheet, p):
        messages.error(request, "You cannot edit this timesheet.")
        return redirect("timesheets")
    entry = get_object_or_404(TimesheetEntry, pk=entry_pk, timesheet=timesheet)
    form = EntryForm(
        request.POST or None,
        request.FILES or None,
        organization=p.organization,
        employee=timesheet.employee,
        instance=entry,
    )
    if request.method == "POST" and form.is_valid():
        updated = form.save(False)
        updated.timesheet = timesheet
        updated.full_clean()
        updated.save()
        _audit(request, "timesheet_entry_updated", updated)
        messages.success(request, "Timesheet entry updated.")
        return redirect("timesheet_detail", pk=pk)
    return render(
        request,
        "erp/timesheet_entry_form.html",
        {
            "form": form,
            "timesheet": timesheet,
            "entry": entry,
            "title": "Edit timesheet entry",
            "submit": "Save changes",
        },
    )


@login_required
@require_POST
def timesheet_entry_delete(request, pk, entry_pk):
    p = _profile(request)
    timesheet = get_object_or_404(Timesheet, pk=pk, organization=p.organization)
    if not _can_edit_timesheet(timesheet, p):
        messages.error(request, "You cannot edit this timesheet.")
        return redirect("timesheets")
    entry = get_object_or_404(TimesheetEntry, pk=entry_pk, timesheet=timesheet)
    _audit(request, "timesheet_entry_deleted", entry)
    entry.delete()
    messages.info(request, "Timesheet entry removed.")
    return redirect("timesheet_detail", pk=pk)


@login_required
@require_POST
@transaction.atomic
def timesheet_action(request, pk):
    p = _profile(request)
    obj = get_object_or_404(
        Timesheet.objects.select_for_update(), pk=pk, organization=p.organization
    )
    action = request.POST.get("action")
    if (
        action == "submit"
        and obj.employee_id == p.id
        and obj.status in ["draft", "rejected"]
    ):
        if not obj.expert_signature:
            messages.error(
                request, "Attach your expert signature before submitting the timesheet."
            )
        elif not obj.entries.exists():
            messages.error(request, "Add at least one entry before submitting.")
        elif obj.entries.filter(hours__lte=0).exists():
            messages.error(
                request, "Enter valid hours for every prefilled task before submitting."
            )
        elif (
            obj.entries.values("date")
            .annotate(total_days=Sum("days_worked"))
            .filter(total_days__gt=1)
            .exists()
        ):
            messages.error(
                request,
                "Daily work exceeds one full day. Adjust overlapping entries before submitting.",
            )
        else:
            obj.status = "submitted"
            obj.requested_approver = obj.employee.manager
            obj.submitted_at = timezone.now()
            obj.save()
            if obj.requested_approver_id:
                approval_task = obj.request_task or Task(
                    organization=p.organization,
                    created_by=request.user,
                    start_date=timezone.localdate(),
                    due_date=timezone.localdate() + timedelta(days=2),
                )
                approval_task.title = (
                    f"Review {obj.period_start:%B %Y} timesheet for {obj.employee}"
                )
                approval_task.description = (
                    f"Review the submitted monthly timesheet for {obj.employee}."
                )
                approval_task.instructions = "Check the recorded activities and signatures, then approve, reject, or move the timesheet to your manager."
                approval_task.assigned_to = obj.requested_approver
                approval_task.department = obj.employee.department
                approval_task.priority = "medium"
                approval_task.status = "assigned"
                approval_task.full_clean()
                approval_task.save()
                if obj.request_task_id != approval_task.id:
                    obj.request_task = approval_task
                    obj.save(update_fields=["request_task", "updated_at"])
                send_task_assignment_email(approval_task)
            recipients = (
                User.objects.filter(profile=obj.requested_approver, is_active=True)
                if obj.requested_approver_id
                else User.objects.filter(
                    profile__organization=p.organization,
                    profile__role="admin",
                    is_active=True,
                )
            )
            Notification.objects.bulk_create(
                [
                    Notification(
                        user=u,
                        title="Timesheet awaiting review",
                        message=str(obj),
                        url=f"/timesheets/{obj.pk}/",
                    )
                    for u in recipients
                ]
            )
            _audit(request, "timesheet_submitted", obj)
    elif (
        action in ["approved", "rejected"]
        and obj.status == "submitted"
        and _can_review_timesheet(obj, p)
    ):
        obj.status = action
        obj.reviewed_by = request.user
        obj.reviewed_at = timezone.now()
        obj.review_notes = request.POST.get("notes", "")
        obj.save()
        Notification.objects.create(
            user=obj.employee.user,
            title=f"Timesheet {action}",
            message=obj.review_notes,
            url=f"/timesheets/{obj.pk}/",
        )
        _audit(request, f"timesheet_{action}", obj)
    else:
        messages.error(request, "That action is not allowed.")
    return redirect("timesheet_detail", pk=pk)


def _can_review_timesheet(timesheet, profile):
    if profile.role == "admin":
        return True
    return profile.role == "manager" and timesheet.requested_approver_id == profile.id


def _manager_can_view_timesheet(timesheet, profile):
    if timesheet.employee_id == profile.id:
        return True
    if timesheet.status == "draft":
        return False
    return (
        timesheet.requested_approver_id == profile.id
        or timesheet.employee_id in _reporting_descendant_ids(profile)
    )


def _can_edit_timesheet(timesheet, profile):
    if timesheet.status == "approved":
        return False
    if timesheet.employee_id == profile.id:
        return timesheet.status in ["draft", "rejected"]
    if profile.role == "admin":
        return timesheet.status in ["draft", "rejected", "submitted"]
    return (
        profile.role == "manager"
        and timesheet.status == "submitted"
        and _manager_can_view_timesheet(timesheet, profile)
    )


def _can_delete_timesheet(timesheet, profile):
    if timesheet.status != "draft":
        return False
    if timesheet.employee_id == profile.id or profile.role == "admin":
        return True
    return False


@login_required
def timesheet_signature_upload(request, pk):
    p = _profile(request)
    obj = get_object_or_404(
        Timesheet,
        pk=pk,
        organization=p.organization,
        employee=p,
        status__in=["draft", "rejected"],
    )
    form = TimesheetSignatureForm(
        request.POST or None, request.FILES or None, instance=obj
    )
    if request.method == "POST" and form.is_valid():
        form.save()
        _audit(request, "timesheet_signature_attached", obj)
        messages.success(request, "Expert signature attached.")
        return redirect("timesheet_detail", pk=pk)
    return render(
        request,
        "erp/timesheet_signature_form.html",
        {"form": form, "timesheet": obj, "title": "Attach expert signature"},
    )


@login_required
@transaction.atomic
def timesheet_review(request, pk):
    p = _profile(request)
    obj = get_object_or_404(
        Timesheet.objects.select_for_update().select_related(
            "employee__user", "employee__manager"
        ),
        pk=pk,
        organization=p.organization,
        status="submitted",
    )
    if not _can_review_timesheet(obj, p):
        messages.error(request, "You cannot review this timesheet.")
        return redirect("timesheets")
    form = TimesheetReviewForm(
        request.POST or None,
        request.FILES or None,
        can_escalate=bool(
            p.role == "manager" and p.manager_id and p.manager.user.is_active
        ),
    )
    if request.method == "POST" and form.is_valid():
        decision = form.cleaned_data["decision"]
        if decision == "escalate":
            next_approver = p.manager
            if not next_approver or not next_approver.user.is_active:
                messages.error(
                    request,
                    "You do not have an active manager available for escalation.",
                )
                return redirect("timesheet_detail", pk=pk)
            obj.requested_approver = next_approver
            obj.review_notes = form.cleaned_data["notes"]
            obj.save(update_fields=["requested_approver", "review_notes", "updated_at"])
            if obj.request_task_id:
                review_task = obj.request_task
                review_task.assigned_to = next_approver
                review_task.department = next_approver.department
                review_task.status = "assigned"
                review_task.priority = "medium"
                review_task.save(
                    update_fields=[
                        "assigned_to",
                        "department",
                        "status",
                        "priority",
                        "updated_at",
                    ]
                )
            Notification.objects.create(
                user=next_approver.user,
                title="Timesheet escalated to you",
                message=str(obj),
                url=f"/timesheets/{obj.pk}/",
            )
            send_task_assignment_email(review_task) if obj.request_task_id else None
            _audit(request, "timesheet_escalated", obj)
            messages.success(request, f"Timesheet moved to {next_approver}.")
            return redirect("timesheet_detail", pk=pk)
        obj.status = decision
        obj.review_notes = form.cleaned_data["notes"]
        obj.reviewed_by = request.user
        obj.reviewed_at = timezone.now()
        if form.cleaned_data.get("manager_signature"):
            obj.manager_signature = form.cleaned_data["manager_signature"]
        if form.cleaned_data.get("consultant_name"):
            obj.consultant_name = form.cleaned_data["consultant_name"]
        if form.cleaned_data.get("consultant_signature"):
            obj.consultant_signature = form.cleaned_data["consultant_signature"]
        obj.save()
        if obj.request_task_id:
            review_task = obj.request_task
            review_task.status = "completed" if obj.status == "approved" else "rejected"
            review_task.save(update_fields=["status", "updated_at"])
        Notification.objects.create(
            user=obj.employee.user,
            title=f"Timesheet {obj.status}",
            message=obj.review_notes,
            url=f"/timesheets/{obj.pk}/",
        )
        _audit(request, f"timesheet_{obj.status}", obj)
        messages.success(request, f"Timesheet {obj.status}.")
        return redirect("timesheet_detail", pk=pk)
    return render(
        request,
        "erp/timesheet_review_form.html",
        {"form": form, "timesheet": obj, "title": "Review monthly timesheet"},
    )


@login_required
def timesheet_signature(request, pk, kind):
    p = _profile(request)
    obj = get_object_or_404(Timesheet, pk=pk, organization=p.organization)
    if p.role == "employee" and obj.employee_id != p.id:
        return HttpResponse(status=404)
    if p.role == "manager" and not _manager_can_view_timesheet(obj, p):
        return HttpResponse(status=404)
    fields = {
        "expert": obj.expert_signature,
        "manager": obj.manager_signature,
        "consultant": obj.consultant_signature,
    }
    signature = fields.get(kind)
    if not signature:
        return HttpResponse(status=404)
    return FileResponse(
        signature.open("rb"),
        content_type=mimetypes.guess_type(signature.name)[0] or "image/png",
    )


@login_required
@roles_allowed("admin", "manager")
def timesheet_request(request):
    p = _profile(request)
    eligible_ids = _reporting_descendant_ids(p) if p.role == "manager" else None
    form = TimesheetRequestForm(
        request.POST or None,
        organization=p.organization,
        manager=p,
        eligible_employee_ids=eligible_ids,
        selected_employee=request.GET.get("employee"),
    )
    if request.method == "POST" and form.is_valid():
        employees = list(form.cleaned_data["employee"])
        months = sorted(int(value) for value in form.cleaned_data["month"])
        year = form.cleaned_data["year"]
        requested = []
        skipped = []
        with transaction.atomic():
            for employee in employees:
                for month in months:
                    period_start = date(year, month, 1)
                    period_end = date(year, month, calendar.monthrange(year, month)[1])
                    sheet, created = Timesheet.objects.get_or_create(
                        organization=p.organization,
                        employee=employee,
                        period_start=period_start,
                        defaults={
                            "period_end": period_end,
                            "country": (
                                p.organization.address.split(",")[-1].strip()
                                if p.organization.address
                                else ""
                            ),
                            "place_of_assignment": "Kenya",
                            "requested_approver": p,
                        },
                    )
                    if sheet.status == "approved" or (
                        sheet.request_task_id
                        and sheet.request_task.status not in ["completed", "rejected"]
                    ):
                        skipped.append(f"{employee} · {calendar.month_name[month]}")
                        continue
                    if created:
                        assigned = Task.objects.filter(
                            organization=p.organization,
                            assigned_to=employee,
                            actual_started_at__date__lte=period_end,
                        ).filter(
                            Q(actual_completed_at__isnull=True)
                            | Q(actual_completed_at__date__gte=period_start)
                        )
                        TimesheetEntry.objects.bulk_create(
                            _prefilled_timesheet_entries(sheet, assigned)
                        )
                    request_task = Task(
                        organization=p.organization,
                        title=f"Complete {calendar.month_name[month]} {year} timesheet",
                        description=f"Complete and submit the monthly timesheet requested by {p}.",
                        instructions=form.cleaned_data["instructions"],
                        assigned_to=employee,
                        department=employee.department,
                        created_by=request.user,
                        priority="medium",
                        status="assigned",
                        start_date=timezone.localdate(),
                        due_date=form.cleaned_data["due_date"],
                    )
                    request_task.full_clean()
                    request_task.save()
                    sheet.request_task = request_task
                    sheet.requested_approver = p
                    sheet.save(
                        update_fields=[
                            "request_task",
                            "requested_approver",
                            "updated_at",
                        ]
                    )
                    Notification.objects.create(
                        user=employee.user,
                        title="Timesheet requested",
                        message=request_task.title,
                        url=f"/timesheets/{sheet.pk}/",
                    )
                    _audit(request, "timesheet_requested", sheet)
                    requested.append(request_task)
        for request_task in requested:
            send_task_assignment_email(request_task)
        if requested:
            messages.success(
                request,
                f"Sent {len(requested)} timesheet request{'s' if len(requested)!=1 else ''} to {len(employees)} employee{'s' if len(employees)!=1 else ''}.",
            )
        if skipped:
            messages.info(
                request,
                f"Skipped {len(skipped)} approved or already-requested timesheet{'s' if len(skipped)!=1 else ''}.",
            )
        return redirect("timesheets")
    return render(
        request,
        "erp/timesheet_request_form.html",
        {"form": form, "title": "Request employee timesheets"},
    )


@login_required
def timesheet_export(request, pk, fmt):
    p = _profile(request)
    obj = get_object_or_404(
        Timesheet.objects.prefetch_related("entries"),
        pk=pk,
        organization=p.organization,
    )
    if p.role == "employee" and obj.employee_id != p.id:
        return redirect("timesheets")
    if p.role == "manager" and not _manager_can_view_timesheet(obj, p):
        return redirect("timesheets")
    if fmt != "pdf":
        return HttpResponse(status=404)

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.setTitle(f"Timesheet {pk}")
    c.setFont("Helvetica-Bold", 18)
    c.drawString(50, 800, p.organization.name)
    c.setFont("Helvetica", 11)
    c.drawString(
        50,
        778,
        f"Timesheet · {obj.employee.user.get_full_name()} · {obj.period_start} to {obj.period_end}",
    )
    y = 740
    for e in obj.entries.all():
        c.drawString(50, y, str(e.date))
        c.drawString(130, y, e.task_performed[:55])
        c.drawRightString(540, y, f"{e.hours} hrs")
        y -= 22
    c.setFont("Helvetica-Bold", 11)
    c.drawString(
        50,
        y - 10,
        f"Total hours: {obj.total_hours}   Status: {obj.get_status_display()}",
    )
    c.save()
    return HttpResponse(
        buf.getvalue(),
        content_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="timesheet-{pk}.pdf"'},
    )


@login_required
def timesheet_export_year(request, pk):
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as ExcelImage
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    p = _profile(request)
    anchor = get_object_or_404(
        Timesheet.objects.select_related("employee__user", "employee__department"),
        pk=pk,
        organization=p.organization,
    )
    if p.role == "employee" and anchor.employee_id != p.id:
        return HttpResponse(status=404)
    if p.role == "manager" and not _manager_can_view_timesheet(anchor, p):
        return HttpResponse(status=404)
    year = anchor.period_start.year
    form = TimesheetExportForm(request.POST or None, year=year)
    if request.method != "POST" or not form.is_valid():
        return render(
            request,
            "erp/timesheet_export_form.html",
            {
                "form": form,
                "timesheet": anchor,
                "year": year,
                "title": "Choose months to export",
            },
        )
    selected_months = sorted({int(month) for month in form.cleaned_data["months"]})
    sheet_query = Timesheet.objects.filter(
        organization=p.organization, employee=anchor.employee, period_start__year=year
    )
    if p.role == "manager" and anchor.employee_id != p.id:
        sheet_query = sheet_query.exclude(status="draft")
    sheets = list(sheet_query.prefetch_related("entries").order_by("period_start"))
    by_month = {sheet.period_start.month: sheet for sheet in sheets}
    metadata = sheets[0] if sheets else anchor
    wb = Workbook()
    thin = Side(style="thin", color="404040")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    label_fill = "E5E7EB"
    header_fill = "4B5563"
    weekend_fill = "F3F4F6"
    signature_buffers = []
    month_names = []
    for export_index, month in enumerate(selected_months):
        ws = wb.active if export_index == 0 else wb.create_sheet()
        ws.title = f"TS - {calendar.month_name[month]} {str(year)[-2:]}"
        month_names.append(ws.title)
        ws.sheet_view.showGridLines = False
        ws.freeze_panes = "A14"
        ws.merge_cells("A1:H1")
        ws["A1"] = "TIME SHEET FOR EXPERTS"
        ws["A1"].font = Font(size=16, bold=True, color="111827")
        ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[1].height = 32
        current = by_month.get(month)
        source = current or metadata
        details = [
            ("SERVICE CONTRACT", source.service_contract if source else ""),
            ("FINANCING", source.financing if source else ""),
            ("Contract number", source.contract_number if source else ""),
            ("Country", source.country if source else ""),
            ("Place of assignment", source.place_of_assignment if source else ""),
            ("Name of the Expert", str(anchor.employee)),
            ("Position", anchor.employee.position),
        ]
        for row, (label, value) in enumerate(details, start=3):
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=2)
            ws.merge_cells(start_row=row, start_column=3, end_row=row, end_column=8)
            ws.cell(row, 1, label)
            ws.cell(row, 3, value or "")
            ws.cell(row, 1).font = Font(bold=True, color="111827")
            ws.cell(row, 1).fill = PatternFill("solid", fgColor=label_fill)
            ws.cell(row, 1).border = border
            ws.cell(row, 3).border = border
            ws.cell(row, 3).alignment = Alignment(wrap_text=True)
        ws.merge_cells("A11:B11")
        ws["A11"] = "Month"
        ws.merge_cells("C11:D11")
        ws["C11"] = calendar.month_name[month]
        ws.merge_cells("E11:F11")
        ws["E11"] = "Year"
        ws.merge_cells("G11:H11")
        ws["G11"] = year
        for cell in ws[11]:
            cell.border = border
            cell.alignment = Alignment(horizontal="center")
            cell.font = Font(bold=True)
        headers = [
            "Date",
            "Days",
            "Days worked",
            "Overnight stay in the country of the duty station",
            "Overnight stay in country outside of duty station",
            "Per Diem requested Overnight stay outside the country",
            "Location of services (country, city & region)",
            "Activity (travel, mission, report, training, workshops, etc)",
        ]
        for col, value in enumerate(headers, start=1):
            cell = ws.cell(13, col, value)
            cell.fill = PatternFill("solid", fgColor=header_fill)
            cell.font = Font(bold=True, color="FFFFFF", size=9)
            cell.alignment = Alignment(
                horizontal="center", vertical="center", wrap_text=True
            )
            cell.border = border
        ws.row_dimensions[13].height = 55
        days_in_month = calendar.monthrange(year, month)[1]
        entries_by_date = {}
        if current:
            for entry in current.entries.all():
                entries_by_date.setdefault(entry.date, []).append(entry)
        for day in range(1, 32):
            row = 13 + day
            if day <= days_in_month:
                current_date = date(year, month, day)
                items = entries_by_date.get(current_date, [])
                ws.cell(row, 1, day)
                ws.cell(row, 2, current_date.strftime("%a"))
                ws.cell(
                    row,
                    3,
                    (
                        float(sum((item.days_worked for item in items), 0))
                        if items
                        else None
                    ),
                )
                ws.cell(
                    row,
                    4,
                    1 if any(item.overnight_duty_station for item in items) else None,
                )
                ws.cell(
                    row,
                    5,
                    (
                        1
                        if any(item.overnight_outside_duty_station for item in items)
                        else None
                    ),
                )
                ws.cell(
                    row,
                    6,
                    1 if any(item.per_diem_requested for item in items) else None,
                )
                ws.cell(
                    row,
                    7,
                    ", ".join(
                        dict.fromkeys(item.location for item in items if item.location)
                    ),
                )
                ws.cell(
                    row,
                    8,
                    "\n".join(
                        f"{item.task_performed}{': '+item.description if item.description else ''}"
                        for item in items
                    ),
                )
                if current_date.weekday() >= 5:
                    for col in range(1, 9):
                        ws.cell(row, col).fill = PatternFill(
                            "solid", fgColor=weekend_fill
                        )
            for col in range(1, 9):
                ws.cell(row, col).border = border
                ws.cell(row, col).alignment = Alignment(vertical="top", wrap_text=True)
            ws.row_dimensions[row].height = 30
        ws["A46"] = "Monthly total"
        ws["C46"] = "=SUM(C14:C44)"
        ws["D46"] = "=SUM(D14:D44)"
        ws["E46"] = "=SUM(E14:E44)"
        ws["F46"] = "=SUM(F14:F44)"
        ws["A47"] = "Initial budget"
        ws["C47"] = (
            float(source.initial_budget_days)
            if source and source.initial_budget_days is not None
            else None
        )
        ws["A48"] = "Days worked since start"
        ws["C48"] = float(
            sum(
                (
                    sheet.total_days_worked
                    for sheet in sheets
                    if sheet.period_start.month <= month
                ),
                0,
            )
        )
        ws["A49"] = "Days remaining"
        ws["C49"] = '=IF(C47="","",C47-C48)'
        ws["A50"] = "Budget consumption rate"
        ws["C50"] = "=IFERROR(C48/C47,0)"
        ws["C50"].number_format = "0.0%"
        for row in range(46, 51):
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=2)
            ws.cell(row, 1).font = Font(bold=True, color="111827")
            ws.cell(row, 1).fill = PatternFill("solid", fgColor=label_fill)
            ws.cell(row, 1).border = border
            ws.cell(row, 3).border = border
        ws.merge_cells("A52:C55")
        ws.merge_cells("D52:H55")
        ws.merge_cells("D58:H61")
        for area in ["A52:C55", "D52:H55", "D58:H61"]:
            for row in ws[area]:
                for cell in row:
                    cell.border = border
        ws.merge_cells("A56:C56")
        ws["A56"] = f"Signature of the Expert ({anchor.employee})"
        ws.merge_cells("D56:H56")
        ws["D56"] = "Signature of the Head of mission / Manager"
        ws.merge_cells("D62:H62")
        ws["D62"] = (
            f"Signature of Consultant ({source.consultant_name if source else ''})"
        )
        ws["A56"].alignment = ws["D56"].alignment = ws["D62"].alignment = Alignment(
            horizontal="center"
        )
        ws["A56"].font = ws["D56"].font = ws["D62"].font = Font(size=8, italic=True)

        def add_signature(field, anchor_cell, width=120, height=45):
            if not field:
                return
            try:
                field.open("rb")
                data = BytesIO(field.read())
                field.close()
                signature_buffers.append(data)
                image = ExcelImage(data)
                image.width = width
                image.height = height
                ws.add_image(image, anchor_cell)
            except Exception:
                pass

        if current:
            add_signature(current.expert_signature, "B52")
            add_signature(current.manager_signature, "F52")
            add_signature(current.consultant_signature, "F58")
        widths = [9, 10, 13, 22, 22, 22, 24, 70]
        for index, width in enumerate(widths, start=1):
            ws.column_dimensions[get_column_letter(index)].width = width
        for row in ws.iter_rows(min_row=1, max_row=62, min_col=1, max_col=8):
            for cell in row:
                cell.border = border
        ws.page_setup.orientation = "landscape"
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = 1
        ws.sheet_properties.pageSetUpPr.fitToPage = True
        ws.print_area = "A1:H62"
        ws.auto_filter.ref = "A13:H44"
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    month_label = (
        calendar.month_name[selected_months[0]].lower()
        if len(selected_months) == 1
        else f"{len(selected_months)}-months"
    )
    filename = f"{anchor.employee.employee_id or anchor.employee.user.username}-timesheet-{year}-{month_label}.xlsx"
    return HttpResponse(
        output.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _leave_balance_data(profile, year):
    balances = list(
        LeaveAllocation.objects.filter(
            organization=profile.organization, employee=profile, year=year
        ).select_related("employee__user", "assigned_by")
    )
    for balance in balances:
        requests = LeaveRequest.objects.filter(
            organization=profile.organization,
            employee=profile,
            leave_type=balance.leave_type,
            start_date__year=year,
        )
        balance.used_days = (
            requests.filter(status="approved").aggregate(total=Sum("days"))["total"]
            or 0
        )
        balance.pending_days = (
            requests.filter(status="pending").aggregate(total=Sum("days"))["total"] or 0
        )
        balance.remaining_days = max(
            0, balance.allocated_days - balance.used_days - balance.pending_days
        )
    return balances


@login_required
def leave_dashboard(request):
    p = _profile(request)
    try:
        year = int(request.GET.get("year", timezone.localdate().year))
    except (TypeError, ValueError):
        year = timezone.localdate().year
    year = max(2000, min(2100, year))
    own_balances = _leave_balance_data(p, year)
    own_requests = LeaveRequest.objects.filter(
        organization=p.organization, employee=p
    ).select_related("requested_approver__user", "reviewed_by")
    team_balances = []
    review_requests = LeaveRequest.objects.none()
    reviewed_requests = LeaveRequest.objects.none()
    if p.role in ["manager", "admin"]:
        employees = Profile.objects.filter(
            organization=p.organization, user__is_active=True
        ).exclude(role="admin")
        if p.role == "manager":
            employees = employees.filter(manager=p)
        for employee in employees.select_related("user"):
            for balance in _leave_balance_data(employee, year):
                team_balances.append(balance)
        review_requests = LeaveRequest.objects.filter(
            organization=p.organization, status="pending"
        ).select_related("employee__user", "requested_approver__user")
        reviewed_requests = LeaveRequest.objects.filter(
            organization=p.organization, status__in=["approved", "rejected"]
        ).select_related("employee__user", "reviewed_by")
        if p.role == "manager":
            review_requests = review_requests.filter(requested_approver=p)
            reviewed_requests = reviewed_requests.filter(requested_approver=p)
    return render(
        request,
        "erp/leave_dashboard.html",
        {
            "year": year,
            "own_balances": own_balances,
            "own_requests": own_requests,
            "team_balances": team_balances,
            "review_requests": review_requests,
            "reviewed_requests": reviewed_requests,
            "can_allocate": p.role in ["manager", "admin"],
        },
    )


@login_required
def leave_allocation_create(request):
    p = _profile(request)
    if p.role not in ["manager", "admin"]:
        messages.error(request, "Only managers can assign leave days.")
        return redirect("leave_dashboard")
    form = LeaveAllocationForm(
        request.POST or None, organization=p.organization, manager=p
    )
    if request.method == "POST" and form.is_valid():
        data = form.cleaned_data
        with transaction.atomic():
            allocation, created = LeaveAllocation.objects.update_or_create(
                organization=p.organization,
                employee=data["employee"],
                leave_type=data["leave_type"],
                year=data["year"],
                defaults={
                    "allocated_days": data["allocated_days"],
                    "assigned_by": request.user,
                    "notes": data["notes"],
                },
            )
            allocation.full_clean()
            allocation.save()
            Notification.objects.create(
                user=allocation.employee.user,
                title="Leave balance updated",
                message=f"You have {allocation.allocated_days} {allocation.get_leave_type_display().lower()} days for {allocation.year}.",
                url="/leave/",
            )
            _audit(
                request,
                "leave_allocation_created" if created else "leave_allocation_updated",
                allocation,
            )
        messages.success(
            request,
            f"Leave days {'assigned' if created else 'updated'} for {allocation.employee}.",
        )
        return redirect(
            "leave_allocation_create"
            if request.POST.get("_add_another")
            else "leave_dashboard"
        )
    return render(
        request,
        "erp/form.html",
        {"form": form, "title": "Assign leave days", "submit": "Save allocation"},
    )


@login_required
def leave_request_create(request):
    p = _profile(request)
    form = LeaveRequestForm(request.POST or None, employee=p)
    if request.method == "POST" and form.is_valid():
        with transaction.atomic():
            leave = form.save(False)
            leave.organization = p.organization
            leave.employee = p
            leave.requested_approver = p.manager
            leave.days = form.calculated_days
            leave.status = "pending"
            leave.full_clean()
            leave.save()
            Notification.objects.create(
                user=p.manager.user,
                title="Leave request awaiting review",
                message=f"{p} requested {leave.days} working days of {leave.get_leave_type_display().lower()}.",
                url="/leave/",
            )
            _audit(request, "leave_requested", leave)
        send_leave_email(leave, "submitted")
        messages.success(request, "Your leave request was sent to your manager.")
        return redirect(
            "leave_request_create"
            if request.POST.get("_add_another")
            else "leave_dashboard"
        )
    return render(
        request,
        "erp/form.html",
        {"form": form, "title": "Request leave", "submit": "Send request"},
    )


@login_required
@transaction.atomic
def leave_request_review(request, pk):
    p = _profile(request)
    requests = (
        LeaveRequest.objects.select_for_update()
        .filter(organization=p.organization)
        .select_related("employee__user", "requested_approver__user")
    )
    if p.role == "manager":
        requests = requests.filter(requested_approver=p)
    elif p.role != "admin":
        requests = requests.none()
    leave = get_object_or_404(requests, pk=pk)
    if leave.status != "pending":
        messages.info(request, "This leave request has already been reviewed.")
        return redirect("leave_dashboard")
    next_approver = p.manager if p.role == "manager" else None
    can_escalate = bool(
        next_approver
        and next_approver.id != p.id
        and next_approver.organization_id == p.organization_id
        and next_approver.role in ["manager", "admin"]
        and next_approver.user.is_active
    )
    form = LeaveReviewForm(request.POST or None, can_escalate=can_escalate)
    if request.method == "POST" and form.is_valid():
        decision = form.cleaned_data["decision"]
        if decision == "escalate":
            if not can_escalate:
                messages.error(
                    request,
                    "You do not have an active manager available for escalation.",
                )
                return redirect("leave_dashboard")
            leave.requested_approver = next_approver
            leave.review_message = form.cleaned_data["message"]
            leave.save(
                update_fields=["requested_approver", "review_message", "updated_at"]
            )
            Notification.objects.create(
                user=next_approver.user,
                title="Leave request escalated to you",
                message=f"{leave.employee} requested {leave.days} working days of {leave.get_leave_type_display().lower()}.",
                url="/leave/",
            )
            Notification.objects.create(
                user=leave.employee.user,
                title="Leave request moved upward",
                message=f"{p} moved your request to {next_approver} for review.",
                url="/leave/",
            )
            _audit(request, "leave_escalated", leave)
            send_leave_email(leave, "submitted")
            messages.success(
                request, f"Leave request moved to {next_approver} for review."
            )
            return redirect("leave_dashboard")
        if decision == "approved":
            allocation = (
                LeaveAllocation.objects.select_for_update()
                .filter(
                    organization=p.organization,
                    employee=leave.employee,
                    leave_type=leave.leave_type,
                    year=leave.start_date.year,
                )
                .first()
            )
            used = (
                LeaveRequest.objects.filter(
                    organization=p.organization,
                    employee=leave.employee,
                    leave_type=leave.leave_type,
                    start_date__year=leave.start_date.year,
                    status="approved",
                )
                .exclude(pk=leave.pk)
                .aggregate(total=Sum("days"))["total"]
                or 0
            )
            if not allocation or used + leave.days > allocation.allocated_days:
                messages.error(
                    request,
                    "This request can no longer be approved because the employee does not have enough leave days available.",
                )
                return redirect("leave_dashboard")
        leave.status = decision
        leave.reviewed_by = request.user
        leave.reviewed_at = timezone.now()
        leave.review_message = form.cleaned_data["message"]
        leave.save(
            update_fields=[
                "status",
                "reviewed_by",
                "reviewed_at",
                "review_message",
                "updated_at",
            ]
        )
        Notification.objects.create(
            user=leave.employee.user,
            title=f"Leave request {leave.get_status_display().lower()}",
            message=leave.review_message
            or f"Your {leave.get_leave_type_display().lower()} request has been {decision}.",
            url="/leave/",
        )
        _audit(request, f"leave_{decision}", leave)
        send_leave_email(leave, "reviewed")
        messages.success(request, f"Leave request {decision}.")
        return redirect("leave_dashboard")
    return render(
        request,
        "erp/leave_review_form.html",
        {"form": form, "leave": leave, "title": "Review leave request"},
    )


def _next_voucher_number(organization):
    year = timezone.localdate().year
    prefix = f"PV-{year}-"
    numbers = PaymentVoucher.objects.filter(
        organization=organization, number__startswith=prefix
    ).values_list("number", flat=True)
    used = []
    for number in numbers:
        try:
            used.append(int(number.rsplit("-", 1)[1]))
        except (TypeError, ValueError):
            pass
    return f"{prefix}{max(used,default=0)+1:04d}"


def _visible_vouchers(profile):
    qs = (
        PaymentVoucher.objects.filter(organization=profile.organization)
        .select_related(
            "department",
            "prepared_by",
            "approved_by",
            "requested_approver__user",
            "approval_task",
        )
        .prefetch_related("lines", "receipts")
    )
    if profile.role == "employee":
        return qs.filter(prepared_by=profile.user)
    if profile.role == "manager":
        return qs.filter(
            Q(prepared_by=profile.user) | Q(requested_approver=profile)
        ).distinct()
    return qs


@login_required
def payment_vouchers(request):
    p = _profile(request)
    qs = _visible_vouchers(p)
    q = request.GET.get("q", "").strip()
    status = request.GET.get("status", "").strip()
    if q:
        qs = qs.filter(
            Q(number__icontains=q)
            | Q(payee__icontains=q)
            | Q(payee_id_number__icontains=q)
            | Q(department__name__icontains=q)
        )
    if status:
        qs = qs.filter(status=status)
    return render(
        request,
        "erp/payment_voucher_list.html",
        {"vouchers": qs, "statuses": PaymentVoucher.STATUSES},
    )


@login_required
def payment_voucher_create(request):
    p = _profile(request)
    approver = p.manager
    if (
        not approver
        or not approver.user.is_active
        or approver.organization_id != p.organization_id
    ):
        messages.error(
            request,
            "Your employee profile needs an active manager before you can create a payment voucher.",
        )
        return redirect("payment_vouchers")
    voucher = PaymentVoucher(
        organization=p.organization,
        prepared_by=request.user,
        requested_approver=approver,
    )
    form = PaymentVoucherForm(
        request.POST or None,
        request.FILES or None,
        organization=p.organization,
        instance=voucher,
    )
    formset = PaymentVoucherLineFormSet(
        request.POST or None, instance=voucher, prefix="lines"
    )
    if request.method == "POST" and form.is_valid() and formset.is_valid():
        with transaction.atomic():
            Organization.objects.select_for_update().get(pk=p.organization_id)
            voucher = form.save(False)
            voucher.organization = p.organization
            voucher.prepared_by = request.user
            voucher.number = _next_voucher_number(p.organization)
            voucher.status = "submitted"
            voucher.full_clean()
            voucher.save()
            formset.instance = voucher
            formset.save()
            for upload in form.cleaned_data.get("receipts", []):
                PaymentVoucherReceipt.objects.create(
                    voucher=voucher,
                    file=upload,
                    original_name=upload.name[:255],
                    uploaded_by=request.user,
                )
            review_task = Task(
                organization=p.organization,
                title=f"Approve payment voucher {voucher.number}",
                description=f"Review the payment voucher for {voucher.payee}, totaling KSh {voucher.total_amount:,.2f}.",
                instructions="Open the linked payment voucher, verify its particulars and amounts, then approve or reject it.",
                assigned_to=voucher.requested_approver,
                department=voucher.department,
                created_by=request.user,
                priority="medium",
                status="assigned",
                start_date=timezone.localdate(),
                due_date=timezone.localdate() + timedelta(days=2),
            )
            review_task.full_clean()
            review_task.save()
            voucher.approval_task = review_task
            voucher.save(update_fields=["approval_task", "updated_at"])
            Notification.objects.create(
                user=review_task.assigned_to.user,
                title="Payment voucher awaiting approval",
                message=f"{voucher.number} · {voucher.payee}",
                url=f"/payment-vouchers/{voucher.pk}/",
            )
            _audit(request, "payment_voucher_submitted", voucher)
            _audit(request, "task_created", review_task)
        send_task_assignment_email(review_task)
        messages.success(
            request,
            f"Payment voucher {voucher.number} was sent to {voucher.requested_approver} for approval.",
        )
        if request.POST.get("_add_another"):
            return redirect("payment_voucher_create")
        return redirect("payment_voucher_detail", pk=voucher.pk)
    return render(
        request,
        "erp/payment_voucher_form.html",
        {
            "form": form,
            "formset": formset,
            "title": "Create payment voucher",
            "automatic_approver": approver,
        },
    )


@login_required
def payment_voucher_detail(request, pk):
    p = _profile(request)
    voucher = get_object_or_404(_visible_vouchers(p), pk=pk)
    return render(request, "erp/payment_voucher_detail.html", {"voucher": voucher})


@login_required
def payment_voucher_receipt_add(request, pk):
    p = _profile(request)
    voucher = get_object_or_404(
        _visible_vouchers(p),
        pk=pk,
        prepared_by=request.user,
        status__in=["draft", "submitted", "rejected"],
    )
    form = PaymentVoucherReceiptForm(request.POST or None, request.FILES or None)
    if request.method == "POST" and form.is_valid():
        uploads = form.cleaned_data["receipts"]
        for upload in uploads:
            PaymentVoucherReceipt.objects.create(
                voucher=voucher,
                file=upload,
                original_name=upload.name[:255],
                uploaded_by=request.user,
            )
        if (
            voucher.requested_approver
            and voucher.requested_approver.user_id != request.user.id
        ):
            Notification.objects.create(
                user=voucher.requested_approver.user,
                title="Receipts added to payment voucher",
                message=f"{len(uploads)} PDF receipt{'s were' if len(uploads)!=1 else ' was'} added to {voucher.number}.",
                url=f"/payment-vouchers/{voucher.pk}/",
            )
        _audit(request, "payment_voucher_receipts_added", voucher)
        messages.success(
            request,
            f"{len(uploads)} receipt{'s' if len(uploads)!=1 else ''} attached to {voucher.number}.",
        )
        if request.POST.get("_add_another"):
            return redirect("payment_voucher_receipt_add", pk=pk)
        return redirect("payment_voucher_detail", pk=pk)
    return render(
        request,
        "erp/form.html",
        {"form": form, "title": "Attach PDF receipts", "submit": "Attach receipts"},
    )


@login_required
def payment_voucher_receipt(request, pk, receipt_pk):
    p = _profile(request)
    voucher = get_object_or_404(_visible_vouchers(p), pk=pk)
    receipt = get_object_or_404(PaymentVoucherReceipt, pk=receipt_pk, voucher=voucher)
    return FileResponse(
        receipt.file.open("rb"),
        content_type="application/pdf",
        as_attachment=False,
        filename=receipt.original_name,
    )


@login_required
@transaction.atomic
def payment_voucher_action(request, pk, action):
    p = _profile(request)
    voucher = get_object_or_404(_visible_vouchers(p).select_for_update(), pk=pk)
    allowed = {
        "submit": (["draft", "rejected"], "submitted"),
        "approve": (["submitted"], "approved"),
        "reject": (["submitted"], "rejected"),
        "escalate": (["submitted"], "submitted"),
        "paid": (["approved"], "paid"),
    }
    reviewer_action = action in ["approve", "reject", "escalate"]
    if (
        action not in allowed
        or (reviewer_action and voucher.requested_approver_id != p.id)
        or (action == "submit" and voucher.prepared_by_id != request.user.id)
        or (action == "paid" and p.role != "admin")
    ):
        messages.error(request, "That voucher action is not allowed.")
        return redirect("payment_voucher_detail", pk=pk)
    form = PaymentVoucherActionForm(request.POST or None, action=action)
    if request.method == "POST" and form.is_valid():
        valid_from, target = allowed[action]
        if voucher.status not in valid_from:
            messages.error(
                request, "The voucher status changed before this action was completed."
            )
        elif action == "escalate":
            next_approver = p.manager
            if (
                not next_approver
                or next_approver.id == p.id
                or next_approver.organization_id != p.organization_id
                or next_approver.role not in ["manager", "admin"]
                or not next_approver.user.is_active
            ):
                messages.error(
                    request,
                    "You do not have an active manager available for escalation.",
                )
            else:
                voucher.requested_approver = next_approver
                voucher.review_notes = form.cleaned_data.get("notes", "")
                voucher.save(
                    update_fields=["requested_approver", "review_notes", "updated_at"]
                )
                review_task = voucher.approval_task
                if review_task:
                    review_task.assigned_to = next_approver
                    review_task.status = "assigned"
                    review_task.start_date = timezone.localdate()
                    review_task.due_date = timezone.localdate() + timedelta(days=2)
                    review_task.save(
                        update_fields=[
                            "assigned_to",
                            "status",
                            "start_date",
                            "due_date",
                            "updated_at",
                        ]
                    )
                Notification.objects.create(
                    user=next_approver.user,
                    title="Payment voucher escalated to you",
                    message=f"{voucher.number} · {voucher.payee}",
                    url=f"/payment-vouchers/{voucher.pk}/",
                )
                if voucher.prepared_by_id != request.user.id:
                    Notification.objects.create(
                        user=voucher.prepared_by,
                        title="Payment voucher escalated",
                        message=f"{p} moved {voucher.number} to {next_approver}.",
                        url=f"/payment-vouchers/{voucher.pk}/",
                    )
                _audit(request, "payment_voucher_escalated", voucher)
                if review_task:
                    send_task_assignment_email(review_task)
                messages.success(
                    request,
                    f"Voucher {voucher.number} was moved to {next_approver} for approval.",
                )
                return redirect("payment_vouchers")
        else:
            voucher.status = target
            voucher.review_notes = form.cleaned_data.get("notes", "")
            fields = ["status", "review_notes", "updated_at"]
            if action == "approve":
                voucher.approved_by = request.user
                voucher.approved_at = timezone.now()
                fields += ["approved_by", "approved_at"]
            if action == "reject":
                voucher.approved_by = None
                voucher.approved_at = None
                fields += ["approved_by", "approved_at"]
            if action == "paid":
                voucher.payment_received_by = form.cleaned_data["payment_received_by"]
                voucher.paid_at = timezone.now()
                fields += ["payment_received_by", "paid_at"]
            voucher.save(update_fields=fields)
            _audit(request, f"payment_voucher_{target}", voucher)
            if reviewer_action and voucher.approval_task_id:
                review_task = voucher.approval_task
                review_task.status = "completed"
                review_task.save(update_fields=["status", "updated_at"])
            if voucher.prepared_by_id != request.user.id:
                Notification.objects.create(
                    user=voucher.prepared_by,
                    title=f"Voucher {target}",
                    message=f"{voucher.number} for {voucher.payee} is now {target}.",
                    url=f"/payment-vouchers/{voucher.pk}/",
                )
            messages.success(request, f"Voucher {voucher.number} marked {target}.")
        return redirect("payment_voucher_detail", pk=pk)
    return render(
        request,
        "erp/payment_voucher_action_form.html",
        {
            "form": form,
            "voucher": voucher,
            "action": action,
            "title": f"{action.title()} voucher",
        },
    )


@login_required
def payment_voucher_pdf(request, pk):
    p = _profile(request)
    voucher = get_object_or_404(
        _visible_vouchers(p).select_related("organization"), pk=pk
    )
    buf = BytesIO()
    pdf = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    logo_path = (
        voucher.organization.logo.path
        if voucher.organization.logo
        else str(settings.BASE_DIR / "static" / "images" / "powerpaylogo.png")
    )
    try:
        pdf.drawImage(
            ImageReader(logo_path),
            50,
            height - 105,
            width=150,
            height=55,
            preserveAspectRatio=True,
            mask="auto",
            anchor="w",
        )
    except Exception:
        pdf.setFont("Helvetica-Bold", 22)
        pdf.setFillColorRGB(0.02, 0.6, 0.35)
        pdf.drawString(50, height - 78, "PowerPay")
    pdf.setFillColorRGB(0, 0, 0)
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawCentredString(width * 0.68, height - 77, "PAYMENT VOUCHER")
    pdf.setFont("Helvetica", 8)
    pdf.drawRightString(
        width - 50,
        height - 96,
        f"{voucher.number}  |  {voucher.get_status_display().upper()}",
    )
    labels = [
        ("Date", voucher.date.strftime("%d %B %Y")),
        ("Payee", voucher.payee),
        ("ID No.", voucher.payee_id_number or "—"),
        ("Department", voucher.department.name if voucher.department else "—"),
    ]
    y = height - 150
    pdf.setFont("Helvetica-Bold", 10)
    for label, value in labels:
        pdf.drawString(55, y, f"{label}:")
        pdf.setFont("Helvetica", 10)
        pdf.drawString(145, y, value)
        pdf.line(140, y - 3, width - 55, y - 3)
        pdf.setFont("Helvetica-Bold", 10)
        y -= 24
    x = 50
    table_top = y - 10
    table_width = width - 100
    amount_width = 115
    header_h = 28
    body_h = 235
    pdf.rect(x, table_top - header_h - body_h, table_width, header_h + body_h)
    pdf.line(x, table_top - header_h, x + table_width, table_top - header_h)
    pdf.line(
        x + table_width - amount_width,
        table_top,
        x + table_width - amount_width,
        table_top - header_h - body_h,
    )
    pdf.setFont("Helvetica-Bold", 10)
    pdf.drawString(x + 8, table_top - 19, "Particulars")
    pdf.drawCentredString(
        x + table_width - amount_width / 2, table_top - 19, "Amount (KSh)"
    )
    row_y = table_top - header_h - 18
    pdf.setFont("Helvetica", 9)
    for line in voucher.lines.all():
        wrapped = textwrap.wrap(line.particulars, width=62) or [""]
        for text_line in wrapped[:2]:
            pdf.drawString(x + 8, row_y, text_line)
            row_y -= 12
        pdf.drawRightString(x + table_width - 8, row_y + 12, f"{line.amount:,.2f}")
        row_y -= 8
        if row_y < table_top - header_h - body_h + 32:
            break
    pdf.setFont("Helvetica-Bold", 10)
    pdf.drawRightString(
        x + table_width - amount_width - 8, table_top - header_h - body_h + 12, "TOTAL"
    )
    pdf.drawRightString(
        x + table_width - 8,
        table_top - header_h - body_h + 12,
        f"{voucher.total_amount:,.2f}",
    )
    sign_top = table_top - header_h - body_h - 32
    sign_w = table_width / 3
    pdf.rect(x, sign_top - 78, table_width, 78)
    pdf.line(x + sign_w, sign_top, x + sign_w, sign_top - 78)
    pdf.line(x + sign_w * 2, sign_top, x + sign_w * 2, sign_top - 78)
    pdf.line(x, sign_top - 26, x + table_width, sign_top - 26)
    pdf.setFont("Helvetica-Bold", 9)
    pdf.drawString(x + 7, sign_top - 18, "Prepared by:")
    pdf.drawString(x + sign_w + 7, sign_top - 18, "Approved by:")
    pdf.drawString(x + sign_w * 2 + 7, sign_top - 18, "Payment Received by:")
    pdf.setFont("Helvetica", 8)
    pdf.drawString(
        x + 7,
        sign_top - 47,
        voucher.prepared_by.get_full_name() or voucher.prepared_by.username,
    )
    pdf.drawString(
        x + sign_w + 7,
        sign_top - 47,
        (
            (voucher.approved_by.get_full_name() or voucher.approved_by.username)
            if voucher.approved_by
            else ""
        ),
    )
    pdf.drawString(x + sign_w * 2 + 7, sign_top - 47, voucher.payment_received_by)
    pdf.setFont("Helvetica", 7)
    pdf.setFillColorRGB(0.35, 0.35, 0.35)
    pdf.drawCentredString(
        width / 2,
        40,
        f"Generated by PowerpayERP · {voucher.organization.name} · {timezone.localtime().strftime('%d %b %Y %H:%M')}",
    )
    pdf.save()
    response = HttpResponse(buf.getvalue(), content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="{voucher.number}.pdf"'
    return response


def _next_item_request_number(organization):
    year = timezone.localdate().year
    prefix = f"IR-{year}-"
    used = []
    for number in ItemRequest.objects.filter(
        organization=organization, number__startswith=prefix
    ).values_list("number", flat=True):
        try:
            used.append(int(number.rsplit("-", 1)[1]))
        except (TypeError, ValueError):
            pass
    return f"{prefix}{max(used,default=0)+1:04d}"


def _visible_item_requests(profile):
    qs = (
        ItemRequest.objects.filter(organization=profile.organization)
        .select_related(
            "requested_by__profile",
            "department",
            "requested_approver__user",
            "approved_by",
            "approval_task",
        )
        .prefetch_related("lines")
    )
    if profile.role == "employee":
        return qs.filter(requested_by=profile.user)
    if profile.role == "manager":
        return qs.filter(
            Q(requested_by=profile.user) | Q(requested_approver=profile)
        ).distinct()
    return qs


@login_required
def item_requests(request):
    p = _profile(request)
    qs = _visible_item_requests(p)
    q = request.GET.get("q", "").strip()
    status = request.GET.get("status", "").strip()
    if q:
        qs = qs.filter(
            Q(number__icontains=q)
            | Q(purpose__icontains=q)
            | Q(lines__item__icontains=q)
        ).distinct()
    if status:
        qs = qs.filter(status=status)
    return render(
        request,
        "erp/item_request_list.html",
        {"item_requests": qs, "statuses": ItemRequest.STATUSES},
    )


@login_required
def item_request_create(request):
    p = _profile(request)
    approver = p.manager
    if (
        not approver
        or not approver.user.is_active
        or approver.organization_id != p.organization_id
    ):
        messages.error(
            request,
            "Your employee profile needs an active manager before you can request items.",
        )
        return redirect("item_requests")
    item_request = ItemRequest(
        organization=p.organization,
        requested_by=request.user,
        requested_approver=approver,
    )
    form = ItemRequestForm(
        request.POST or None,
        organization=p.organization,
        profile=p,
        instance=item_request,
    )
    formset = ItemRequestLineFormSet(
        request.POST or None, instance=item_request, prefix="lines"
    )
    if request.method == "POST" and form.is_valid() and formset.is_valid():
        with transaction.atomic():
            Organization.objects.select_for_update().get(pk=p.organization_id)
            item_request = form.save(False)
            item_request.organization = p.organization
            item_request.requested_by = request.user
            item_request.requested_approver = approver
            item_request.number = _next_item_request_number(p.organization)
            item_request.status = "submitted"
            item_request.full_clean()
            item_request.save()
            formset.instance = item_request
            formset.save()
            task = Task(
                organization=p.organization,
                title=f"Approve item request {item_request.number}",
                description=f"Review {item_request.lines.count()} requested item(s), estimated at KSh {item_request.total_estimated_cost:,.2f}.",
                instructions="Review quantities, estimated prices and source links, then approve, reject, or move the request to your manager.",
                assigned_to=approver,
                department=item_request.department,
                created_by=request.user,
                priority="medium",
                status="assigned",
                start_date=timezone.localdate(),
                due_date=min(
                    item_request.needed_by, timezone.localdate() + timedelta(days=2)
                ),
            )
            task.full_clean()
            task.save()
            item_request.approval_task = task
            item_request.save(update_fields=["approval_task", "updated_at"])
            Notification.objects.create(
                user=approver.user,
                title="Item request awaiting approval",
                message=f"{item_request.number} · KSh {item_request.total_estimated_cost:,.2f}",
                url=f"/item-requests/{item_request.pk}/",
            )
            _audit(request, "item_request_submitted", item_request)
            _audit(request, "task_created", task)
        send_task_assignment_email(task)
        messages.success(
            request, f"Item request {item_request.number} was sent to {approver}."
        )
        if request.POST.get("_add_another"):
            return redirect("item_request_create")
        return redirect("item_request_detail", pk=item_request.pk)
    return render(
        request,
        "erp/item_request_form.html",
        {"form": form, "formset": formset, "automatic_approver": approver},
    )


@login_required
def item_request_detail(request, pk):
    p = _profile(request)
    obj = get_object_or_404(_visible_item_requests(p), pk=pk)
    return render(request, "erp/item_request_detail.html", {"item_request": obj})


@login_required
@transaction.atomic
def item_request_action(request, pk, action):
    p = _profile(request)
    obj = get_object_or_404(_visible_item_requests(p).select_for_update(), pk=pk)
    if (
        action not in ["approve", "reject", "escalate"]
        or obj.status != "submitted"
        or (p.role != "admin" and obj.requested_approver_id != p.id)
    ):
        messages.error(request, "That action is not allowed.")
        return redirect("item_request_detail", pk=pk)
    form = ItemRequestActionForm(request.POST or None)
    if request.method == "POST" and form.is_valid():
        notes = form.cleaned_data["notes"]
        if action == "escalate":
            next_approver = p.manager
            if not next_approver or not next_approver.user.is_active:
                messages.error(
                    request,
                    "You do not have an active manager available for escalation.",
                )
                return redirect("item_request_detail", pk=pk)
            obj.requested_approver = next_approver
            obj.review_notes = notes
            obj.save(update_fields=["requested_approver", "review_notes", "updated_at"])
            task = obj.approval_task
            task.assigned_to = next_approver
            task.department = next_approver.department
            task.priority = "medium"
            task.status = "assigned"
            task.save(
                update_fields=[
                    "assigned_to",
                    "department",
                    "priority",
                    "status",
                    "updated_at",
                ]
            )
            Notification.objects.create(
                user=next_approver.user,
                title="Item request escalated to you",
                message=obj.number,
                url=f"/item-requests/{obj.pk}/",
            )
            send_task_assignment_email(task)
            _audit(request, "item_request_escalated", obj)
            messages.success(request, f"Item request moved to {next_approver}.")
        else:
            obj.status = "approved" if action == "approve" else "rejected"
            obj.review_notes = notes
            obj.approved_by = request.user if action == "approve" else None
            obj.approved_at = timezone.now() if action == "approve" else None
            obj.save(
                update_fields=[
                    "status",
                    "review_notes",
                    "approved_by",
                    "approved_at",
                    "updated_at",
                ]
            )
            if obj.approval_task_id:
                approval_task = obj.approval_task
                approval_task.status = "completed"
                approval_task.save(update_fields=["status", "updated_at"])
            Notification.objects.create(
                user=obj.requested_by,
                title=f"Item request {obj.status}",
                message=notes or obj.number,
                url=f"/item-requests/{obj.pk}/",
            )
            _audit(request, f"item_request_{obj.status}", obj)
            messages.success(request, f"Item request {obj.status}.")
        return redirect("item_request_detail", pk=pk)
    labels = {
        "approve": "Approve item request",
        "reject": "Reject item request",
        "escalate": "Move item request to my manager",
    }
    return render(
        request,
        "erp/item_request_action_form.html",
        {"form": form, "item_request": obj, "title": labels[action], "action": action},
    )


@login_required
def item_request_export(request, pk, fmt):
    p = _profile(request)
    obj = get_object_or_404(_visible_item_requests(p), pk=pk)
    lines = list(obj.lines.all())
    if fmt == "xlsx":
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

        wb = Workbook()
        ws = wb.active
        ws.title = "Consumable supplies"
        ws.sheet_view.showGridLines = False
        ws.merge_cells("A1:E1")
        ws["A1"] = f"ITEM REQUEST {obj.number}"
        ws["A1"].font = Font(size=18, bold=True, color="FFFFFF")
        ws["A1"].fill = PatternFill("solid", fgColor="087F5B")
        ws["A1"].alignment = Alignment(horizontal="center")
        ws.row_dimensions[1].height = 30
        metadata = [
            (
                "Requested by",
                obj.requested_by.get_full_name() or obj.requested_by.username,
            ),
            ("Department", str(obj.department or "")),
            ("Purpose", obj.purpose),
            ("Needed by", obj.needed_by),
            ("Delivery location", obj.delivery_location),
            ("Status", obj.get_status_display()),
            ("Current approver", str(obj.requested_approver or "")),
        ]
        for row, (label, value) in enumerate(metadata, 3):
            ws.cell(row, 1, label).font = Font(bold=True, color="087F5B")
            ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=5)
            ws.cell(row, 2, value)
            ws.cell(row, 2).alignment = Alignment(wrap_text=True)
        header_row = 11
        headers = ["ITEMS", "QUANTITY", "ESTIMATED PRICE (KSh)", "SOURCE LINK", "NOTES"]
        for col, value in enumerate(headers, 1):
            cell = ws.cell(header_row, col, value)
            cell.fill = PatternFill("solid", fgColor="F97316")
            cell.font = Font(bold=True, color="FFFFFF")
            cell.alignment = Alignment(horizontal="center", wrap_text=True)
        for row, line in enumerate(lines, header_row + 1):
            ws.cell(row, 1, line.item)
            ws.cell(row, 2, line.quantity)
            ws.cell(row, 3, float(line.estimated_cost))
            ws.cell(row, 4, line.source_link)
            ws.cell(row, 5, line.notes)
        total_row = header_row + len(lines) + 1
        ws.cell(total_row, 1, "TOTAL").font = Font(bold=True, color="087F5B")
        ws.cell(total_row, 3, f"=SUM(C{header_row+1}:C{total_row-1})")
        ws.cell(total_row, 3).font = Font(bold=True)
        ws.cell(total_row, 3).number_format = '"KSh" #,##0.00'
        thin = Side(style="thin", color="D1D5DB")
        for row in ws.iter_rows(
            min_row=header_row, max_row=total_row, min_col=1, max_col=5
        ):
            for cell in row:
                cell.border = Border(bottom=thin)
                cell.alignment = Alignment(vertical="top", wrap_text=True)
        ws.column_dimensions["A"].width = 28
        ws.column_dimensions["B"].width = 16
        ws.column_dimensions["C"].width = 22
        ws.column_dimensions["D"].width = 48
        ws.column_dimensions["E"].width = 32
        ws.freeze_panes = f"A{header_row+1}"
        ws.auto_filter.ref = f"A{header_row}:E{total_row-1}"
        ws.page_setup.orientation = "landscape"
        ws.print_area = f"A1:E{total_row}"
        buf = BytesIO()
        wb.save(buf)
        content_type = (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        extension = "xlsx"
    elif fmt == "docx":
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor

        document = Document()
        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = (
            section.page_height,
            section.page_width,
        )
        section.top_margin = section.bottom_margin = Inches(0.65)
        section.left_margin = section.right_margin = Inches(0.7)
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f"ITEM REQUEST {obj.number}")
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(8, 127, 91)
        for label, value in [
            (
                "Requested by",
                obj.requested_by.get_full_name() or obj.requested_by.username,
            ),
            ("Department", obj.department or "-"),
            ("Purpose", obj.purpose),
            ("Needed by", obj.needed_by.strftime("%d %b %Y")),
            ("Delivery location", obj.delivery_location or "-"),
            ("Status", obj.get_status_display()),
            ("Current approver", obj.requested_approver or "-"),
        ]:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.add_run(f"{label}: ").bold = True
            paragraph.add_run(str(value))
        table = document.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for cell, label in zip(
            table.rows[0].cells,
            ["Item", "Quantity", "Estimated price (KSh)", "Source link", "Notes"],
        ):
            cell.text = label
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "087F5B")
            cell._tc.get_or_add_tcPr().append(shading)
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
        for line in lines:
            cells = table.add_row().cells
            for cell, value in zip(
                cells,
                [
                    line.item,
                    line.quantity,
                    f"{line.estimated_cost:,.2f}",
                    line.source_link,
                    line.notes,
                ],
            ):
                cell.text = str(value)
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(8)
        total = document.add_paragraph()
        total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        total.add_run(
            f"Total estimated cost: KSh {obj.total_estimated_cost:,.2f}"
        ).bold = True
        if obj.review_notes:
            document.add_heading("Review notes", level=2)
            document.add_paragraph(obj.review_notes)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run(
            f"Generated by PowerpayERP - {obj.organization.name}"
        ).font.size = Pt(8)
        buf = BytesIO()
        document.save(buf)
        content_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        extension = "docx"
    elif fmt == "pdf":
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            Table,
            TableStyle,
        )

        buf = BytesIO()
        styles = getSampleStyleSheet()
        doc = SimpleDocTemplate(
            buf,
            pagesize=landscape(A4),
            rightMargin=14 * mm,
            leftMargin=14 * mm,
            topMargin=12 * mm,
            bottomMargin=12 * mm,
        )
        story = [
            Paragraph(f"ITEM REQUEST {obj.number}", styles["Title"]),
            Spacer(1, 5 * mm),
            Paragraph(
                f"<b>Requested by:</b> {obj.requested_by.get_full_name() or obj.requested_by.username} &nbsp;&nbsp; <b>Department:</b> {obj.department or '-'} &nbsp;&nbsp; <b>Needed by:</b> {obj.needed_by:%d %b %Y}",
                styles["BodyText"],
            ),
            Paragraph(f"<b>Purpose:</b> {obj.purpose}", styles["BodyText"]),
            Spacer(1, 4 * mm),
        ]
        header_style = ParagraphStyle(
            "RequestHeader",
            parent=styles["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=8,
            leading=9,
            textColor=colors.white,
        )
        data = (
            [
                [
                    Paragraph(label, header_style)
                    for label in [
                        "ITEMS",
                        "QUANTITY",
                        "ESTIMATED PRICE (KSh)",
                        "SOURCE LINK",
                        "NOTES",
                    ]
                ]
            ]
            + [
                [
                    Paragraph(line.item, styles["BodyText"]),
                    line.quantity,
                    f"{line.estimated_cost:,.2f}",
                    Paragraph(line.source_link or "-", styles["BodyText"]),
                    Paragraph(line.notes or "-", styles["BodyText"]),
                ]
                for line in lines
            ]
            + [["TOTAL", "", f"{obj.total_estimated_cost:,.2f}", "", ""]]
        )
        table = Table(
            data, colWidths=[48 * mm, 28 * mm, 38 * mm, 90 * mm, 55 * mm], repeatRows=1
        )
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#087F5B")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#FFF7ED")),
                    ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D1D5DB")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        story.append(table)
        story.extend(
            [
                Spacer(1, 4 * mm),
                Paragraph(
                    f"<b>Status:</b> {obj.get_status_display()} &nbsp;&nbsp; <b>Current approver:</b> {obj.requested_approver or '-'}",
                    styles["BodyText"],
                ),
            ]
        )
        doc.build(story)
        extension = "pdf"
        content_type = "application/pdf"
    else:
        return HttpResponse(status=404)
    response = HttpResponse(buf.getvalue(), content_type=content_type)
    response["Content-Disposition"] = f'attachment; filename="{obj.number}.{extension}"'
    return response


@login_required
@roles_allowed("admin", "manager")
def employees(request):
    p = _profile(request)
    qs = Profile.objects.filter(organization=p.organization).select_related(
        "user", "department", "manager__user"
    )
    if p.role == "manager":
        descendant_ids = _reporting_descendant_ids(p)
        qs = qs.filter(Q(pk=p.pk) | Q(pk__in=descendant_ids)).distinct()
    employees_list = list(
        qs.order_by("department__name", "user__first_name", "user__last_name")
    )
    visible_ids = {employee.id for employee in employees_list}
    eligible_timesheet_ids = (
        set(_reporting_descendant_ids(p))
        if p.role == "manager"
        else {employee.id for employee in employees_list if employee.id != p.id}
    )
    for employee in employees_list:
        employee.can_request_timesheet = employee.id in eligible_timesheet_ids
    hierarchy_groups = []
    departments = list(
        Department.objects.filter(organization=p.organization).order_by("name")
    )
    for department in departments + [None]:
        members = [
            employee
            for employee in employees_list
            if employee.department_id == (department.id if department else None)
        ]
        if not members:
            continue
        children = {}
        for employee in members:
            children.setdefault(employee.manager_id, []).append(employee)
        ordered = []
        visited = set()

        def add_branch(employee, depth):
            if employee.id in visited:
                return
            visited.add(employee.id)
            ordered.append({"employee": employee, "depth": depth})
            for report in children.get(employee.id, []):
                add_branch(report, depth + 1)

        roots = [
            employee
            for employee in members
            if employee.manager_id not in visible_ids
            or employee.manager_id not in {member.id for member in members}
        ]
        for employee in roots:
            add_branch(employee, 0)
        for employee in members:
            add_branch(employee, 0)
        hierarchy_groups.append({"department": department, "nodes": ordered})
    pending = (
        UserInvite.objects.filter(organization=p.organization, is_used=False)
        .select_related("department", "manager__user")
        .order_by("-created_at")
        if p.role == "admin"
        else UserInvite.objects.none()
    )
    return render(
        request,
        "erp/employee_list.html",
        {
            "employees": employees_list,
            "pending_invites": pending,
            "hierarchy_groups": hierarchy_groups,
        },
    )


@login_required
@roles_allowed("admin")
@transaction.atomic
def employee_create(request):
    p = _profile(request)
    form = EmployeeForm(request.POST or None, organization=p.organization)
    if request.method == "POST" and form.is_valid():
        Organization.objects.select_for_update().get(pk=p.organization_id)
        obj = form.save(False)
        obj.organization = p.organization
        obj.invited_by = request.user
        obj.save()
        send_invite_email(obj)
        _audit(request, "employee_invited", obj)
        messages.success(
            request,
            f"Employee {obj.employee_id} created and an activation invitation was sent to {obj.email}.",
        )
        return redirect(
            "employee_create" if request.POST.get("_add_another") else "employees"
        )
    return render(
        request,
        "erp/form.html",
        {"form": form, "title": "Add employee", "submit": "Create and invite"},
    )


@login_required
@roles_allowed("admin")
def departments(request):
    p = _profile(request)
    form = DepartmentForm(request.POST or None)
    if request.method == "POST" and form.is_valid():
        obj = form.save(False)
        obj.organization = p.organization
        obj.save()
        _audit(request, "department_created", obj)
        messages.success(request, "Department created.")
        return redirect("departments")
    return render(
        request,
        "erp/departments.html",
        {
            "departments": Department.objects.filter(
                organization=p.organization
            ).annotate(member_count=Count("members")),
            "form": form,
        },
    )


@login_required
def documents(request):
    p = _profile(request)
    qs = Document.objects.filter(organization=p.organization).select_related(
        "uploaded_by", "department"
    )
    if p.role == "employee":
        qs = qs.filter(
            Q(visibility="organization")
            | Q(department=p.department, visibility="department")
            | Q(owner=p)
        ).distinct()
    return render(request, "erp/document_list.html", {"documents": qs})


def _visible_documents(profile):
    queryset = Document.objects.filter(organization=profile.organization)
    if profile.role == "employee":
        queryset = queryset.filter(
            Q(visibility="organization")
            | Q(department=profile.department, visibility="department")
            | Q(owner=profile)
        )
    return queryset.distinct()


@login_required
def document_download(request, pk):
    if request.user.is_superuser:
        document = get_object_or_404(Document, pk=pk)
        return FileResponse(
            document.file.open("rb"),
            as_attachment=True,
            filename=document.file.name.rsplit("/", 1)[-1],
        )
    profile = _profile(request)
    document = get_object_or_404(_visible_documents(profile), pk=pk)
    return FileResponse(
        document.file.open("rb"),
        as_attachment=True,
        filename=document.file.name.rsplit("/", 1)[-1],
    )


@login_required
def document_upload(request):
    p = _profile(request)
    form = DocumentForm(
        request.POST or None, request.FILES or None, organization=p.organization
    )
    if request.method == "POST" and form.is_valid():
        obj = form.save(False)
        obj.organization = p.organization
        obj.uploaded_by = request.user
        obj.save()
        _audit(request, "document_uploaded", obj)
        messages.success(request, "Document uploaded.")
        return redirect(
            "document_upload" if request.POST.get("_add_another") else "documents"
        )
    return render(
        request,
        "erp/form.html",
        {"form": form, "title": "Upload document", "submit": "Upload"},
    )


@login_required
@roles_allowed("admin")
def templates(request):
    p = _profile(request)
    form = TemplateForm(
        request.POST or None, request.FILES or None, organization=p.organization
    )
    if request.method == "POST" and form.is_valid():
        obj = form.save(False)
        obj.organization = p.organization
        obj.uploaded_by = request.user
        obj.save()
        _audit(request, "template_uploaded", obj)
        messages.success(request, "Template uploaded.")
        return redirect("templates")
    return render(
        request,
        "erp/templates.html",
        {
            "templates": DocumentTemplate.objects.filter(organization=p.organization),
            "form": form,
        },
    )


@login_required
@roles_allowed("admin")
def template_download(request, pk):
    profile = _profile(request)
    template = get_object_or_404(
        DocumentTemplate,
        pk=pk,
        organization=profile.organization,
    )
    return FileResponse(
        template.file.open("rb"),
        as_attachment=True,
        filename=template.file.name.rsplit("/", 1)[-1],
    )


@login_required
def notifications(request):
    return render(
        request,
        "erp/notifications.html",
        {
            "notifications": request.user.notifications.all(),
            "unread_exists": request.user.notifications.filter(is_read=False).exists(),
        },
    )


@login_required
@require_POST
def notification_read(request, pk):
    notification = get_object_or_404(Notification, pk=pk, user=request.user)
    if not notification.is_read:
        notification.is_read = True
        notification.save(update_fields=["is_read", "updated_at"])
    return JsonResponse(
        {"ok": True, "unread": request.user.notifications.filter(is_read=False).count()}
    )


@login_required
@require_POST
def notifications_read_all(request):
    request.user.notifications.filter(is_read=False).update(
        is_read=True, updated_at=timezone.now()
    )
    return JsonResponse({"ok": True, "unread": 0})


@login_required
@roles_allowed("admin")
def audit_logs(request):
    return render(
        request,
        "erp/audit.html",
        {
            "logs": AuditLog.objects.filter(
                organization=_profile(request).organization
            ).select_related("actor")[:200]
        },
    )


def _superuser_required(view):
    return user_passes_test(lambda user: user.is_superuser, login_url="login")(view)


@login_required
@_superuser_required
def platform_organizations(request):
    q = request.GET.get("q", "").strip()
    organizations = Organization.objects.annotate(
        user_count=Count("profiles", distinct=True),
        department_count=Count("departments", distinct=True),
        task_count=Count("task", distinct=True),
    ).order_by("name")
    if q:
        organizations = organizations.filter(
            Q(name__icontains=q)
            | Q(business_email__icontains=q)
            | Q(industry__icontains=q)
        )
    return render(
        request, "erp/platform_organizations.html", {"organizations": organizations}
    )


@login_required
@_superuser_required
def platform_organization_detail(request, pk):
    org = get_object_or_404(Organization, pk=pk)
    return render(
        request,
        "erp/platform_organization_detail.html",
        {
            "organization": org,
            "stats": [
                ("Employees", Profile.objects.filter(organization=org).count()),
                ("Departments", Department.objects.filter(organization=org).count()),
                ("Tasks", Task.objects.filter(organization=org).count()),
                ("Timesheets", Timesheet.objects.filter(organization=org).count()),
                (
                    "Leave requests",
                    LeaveRequest.objects.filter(organization=org).count(),
                ),
                ("Vouchers", PaymentVoucher.objects.filter(organization=org).count()),
                ("Documents", Document.objects.filter(organization=org).count()),
            ],
            "employees": Profile.objects.filter(organization=org).select_related(
                "user", "department", "manager__user"
            )[:25],
            "tasks": Task.objects.filter(organization=org).select_related(
                "assigned_to__user", "department"
            )[:20],
            "timesheets": Timesheet.objects.filter(organization=org).select_related(
                "employee__user", "reviewed_by"
            )[:20],
            "leave_requests": LeaveRequest.objects.filter(
                organization=org
            ).select_related(
                "employee__user", "requested_approver__user", "reviewed_by"
            )[
                :20
            ],
            "vouchers": PaymentVoucher.objects.filter(organization=org)
            .select_related("department", "prepared_by", "approved_by")
            .prefetch_related("lines")[:20],
            "documents": Document.objects.filter(organization=org).select_related(
                "uploaded_by", "department"
            )[:20],
            "activity": AuditLog.objects.filter(organization=org).select_related(
                "actor"
            )[:20],
        },
    )


@login_required
@_superuser_required
@require_POST
def platform_organization_status(request, pk):
    org = get_object_or_404(Organization, pk=pk)
    org.is_active = not org.is_active
    org.save(update_fields=["is_active", "updated_at"])
    AuditLog.objects.create(
        organization=org,
        actor=request.user,
        action="organization_activated" if org.is_active else "organization_suspended",
        entity_type="Organization",
        entity_id=str(org.pk),
        ip_address=request.META.get("REMOTE_ADDR"),
    )
    messages.success(
        request, f"{org.name} has been {'activated' if org.is_active else 'suspended'}."
    )
    if request.POST.get("next") == f"/platform/organizations/{org.pk}/":
        return redirect("platform_organization_detail", pk=org.pk)
    return redirect("platform_organizations")


@login_required
@_superuser_required
def platform_users(request):
    q = request.GET.get("q", "").strip()
    users = User.objects.select_related(
        "profile__organization", "profile__department"
    ).order_by("-date_joined")
    if q:
        users = users.filter(
            Q(email__icontains=q)
            | Q(username__icontains=q)
            | Q(first_name__icontains=q)
            | Q(last_name__icontains=q)
            | Q(profile__organization__name__icontains=q)
        )
    return render(request, "erp/platform_users.html", {"platform_users": users})


@login_required
@_superuser_required
@require_POST
def platform_user_status(request, pk):
    user = get_object_or_404(User, pk=pk)
    if user == request.user:
        messages.error(request, "You cannot deactivate your own platform account.")
    else:
        user.is_active = not user.is_active
        user.save(update_fields=["is_active"])
        profile = getattr(user, "profile", None)
        AuditLog.objects.create(
            organization=getattr(profile, "organization", None),
            actor=request.user,
            action="user_activated" if user.is_active else "user_deactivated",
            entity_type="User",
            entity_id=str(user.pk),
        )
        messages.success(
            request,
            f"{user.email or user.username} has been {'activated' if user.is_active else 'deactivated'}.",
        )
    return redirect("platform_users")


@login_required
@_superuser_required
def platform_activity(request):
    organization = request.GET.get("organization", "")
    action = request.GET.get("action", "").strip()
    logs = AuditLog.objects.select_related("actor", "organization")
    if organization:
        logs = logs.filter(organization_id=organization)
    if action:
        logs = logs.filter(action__icontains=action)
    return render(
        request,
        "erp/platform_activity.html",
        {"logs": logs[:500], "organizations": Organization.objects.order_by("name")},
    )
